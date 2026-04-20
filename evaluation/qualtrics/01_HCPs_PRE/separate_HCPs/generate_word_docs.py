#!/usr/bin/env python3
"""
Generate Word documents from LaTeX sources for both
2_cases_per_HCP (HCP_1–5) and 1_case_per_HCP (HCP_01–10).

Reads each main.tex, extracts text, and builds main.docx.
"""

import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── colour constants ──────────────────────────────────────────────────────────
C_DARK_BLUE   = RGBColor(0x1E, 0x3A, 0x5F)
C_PRIMARY     = RGBColor(0x1D, 0x4E, 0xD8)
C_FOREST      = RGBColor(0x04, 0x78, 0x57)
C_PURPLE      = RGBColor(0x6D, 0x28, 0xD9)
C_AMBER       = RGBColor(0xB4, 0x53, 0x09)
C_SLATE       = RGBColor(0x47, 0x55, 0x69)
C_RED         = RGBColor(0xB9, 0x1C, 0x1C)
C_BLACK       = RGBColor(0x00, 0x00, 0x00)


# ── LaTeX → plain text ────────────────────────────────────────────────────────
def latex2txt(s: str) -> str:
    """Strip common LaTeX markup to get plain Unicode text."""
    # Replace common commands
    subs = [
        (r'\\textbf\{([^}]*)\}',       r'\1'),
        (r'\\textit\{([^}]*)\}',       r'\1'),
        (r'\\emph\{([^}]*)\}',         r'\1'),
        (r'\\small',                    ''),
        (r'\\medskip',                  ''),
        (r'\\smallskip',                ''),
        (r'\\bigskip',                  ''),
        (r'\\vspace\{[^}]*\}',         ''),
        (r'\\hspace\{[^}]*\}',         ' '),
        (r'\\hspace\*\{[^}]*\}',       ' '),
        (r'\\noindent',                 ''),
        (r'\\color\{[^}]*\}',          ''),
        (r'\\textcolor\{[^}]*\}\{([^}]*)\}', r'\1'),
        (r'\\enspace',                  ' '),
        (r'\\quad',                     '  '),
        (r'\\hfill',                    '    '),
        (r'\\par\b',                    '\n'),
        (r'\\\\(\[[^\]]*\])?',          '\n'),
        (r'\\newline',                  '\n'),
        (r'\\begin\{itemize\}[^\n]*',   ''),
        (r'\\end\{itemize\}',           ''),
        (r'\\begin\{enumerate\}[^\n]*', ''),
        (r'\\end\{enumerate\}',         ''),
        (r'\\item\s*',                  '• '),
        (r'\$\\rightarrow\$',           '→'),
        (r'\$\\times\$',               '×'),
        (r'\$\\approx\\,',             '≈\u202f'),
        (r'\$\\approx\$',              '≈'),
        (r'\\approx',                   '≈'),
        (r'\$\\square\$',              '☐'),
        (r'\\square',                   '☐'),
        (r'\\checkmark',                '✓'),
        (r'---',                        '—'),
        (r'--',                         '–'),
        (r'``',                         '\u201c'),
        (r"''",                         '\u201d'),
        (r'`',                          '\u2018'),
        (r"'",                          '\u2019'),
        (r'\\texttt\{([^}]*)\}',       r'\1'),
        (r'\\label\{[^}]*\}',          ''),
        (r'\\ref\{[^}]*\}',            '?'),
        (r'\\pageref\{[^}]*\}',        '?'),
        (r'\\ ',                        ' '),
        (r'\\,',                        '\u202f'),
        (r'\\;',                        ' '),
        (r'\\:',                        ' '),
        (r'\\ ',                        ' '),
        (r'\\!',                        ''),
        (r"\\'{([^}])}",               '\\1\u0301'),  # accent
        (r'\\`\{([^}])\}',             '\\1\u0300'),
        (r'\\text\{([^}]*)\}',         r'\1'),
        (r'\\[a-z]+\*?\{([^}]*)\}',    r'\1'),  # generic command with one arg
        (r'\\[a-z]+\*?',               ''),      # bare commands
        (r'\{|\}',                      ''),
        (r'\$',                         ''),
    ]
    for pat, rep in subs:
        s = re.sub(pat, rep, s, flags=re.DOTALL)
    s = re.sub(r'[ \t]+', ' ', s)
    s = re.sub(r'\n{3,}', '\n\n', s)
    return s.strip()


# ── document helpers ──────────────────────────────────────────────────────────
def set_font(run, bold=False, italic=False, size=11, color=None):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1, color=C_DARK_BLUE):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = color
    return p


def add_para(doc, text='', bold=False, italic=False, size=11,
             color=C_BLACK, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_font(run, bold=bold, italic=italic, size=size, color=color)
    return p


def add_box(doc, title, content_lines, title_color=C_DARK_BLUE,
            shade_color="E8F0FE"):
    """Add a bordered, shaded box with title and content."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]

    # Title paragraph
    tp = cell.paragraphs[0]
    tp.paragraph_format.space_before = Pt(2)
    tp.paragraph_format.space_after = Pt(3)
    tr = tp.add_run(title)
    tr.bold = True
    tr.font.color.rgb = title_color
    tr.font.size = Pt(10)

    # Shade cell
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), shade_color)
    tcPr.append(shd)

    # Content
    for line in content_lines:
        if line.strip() == '':
            cell.add_paragraph('')
            continue
        cp = cell.add_paragraph()
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after = Pt(1)
        cr = cp.add_run(line)
        cr.font.size = Pt(9.5)

    doc.add_paragraph('')  # spacing after box


def add_answer_box(doc, label='Uw antwoord', lines=4):
    """Add a yellow-ish answer box with blank write lines."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFFDE7')
    tcPr.append(shd)

    lp = cell.paragraphs[0]
    lr = lp.add_run(label)
    lr.bold = True
    lr.font.size = Pt(10)
    lr.font.color.rgb = C_PURPLE

    for _ in range(lines):
        bp = cell.add_paragraph('_' * 90)
        bp.paragraph_format.space_before = Pt(6)
        bp.paragraph_format.space_after = Pt(2)
        br = bp.runs[0]
        br.font.size = Pt(8)
        br.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    doc.add_paragraph('')


def add_slot_box(doc, label, n):
    """Add a labeled slot box (for symptoom/behandelingsoptie labels)."""
    for i in range(1, n + 1):
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = 'Table Grid'
        c0 = tbl.rows[0].cells[0]
        c1 = tbl.rows[0].cells[1]
        c0.width = Cm(4)
        c1.width = Cm(12)

        p0 = c0.paragraphs[0]
        r0 = p0.add_run(f'{label} {i}  Label:')
        r0.bold = True
        r0.font.size = Pt(9)
        r0.font.color.rgb = C_SLATE

        p1 = c1.paragraphs[0]
        r1 = p1.add_run('_' * 70)
        r1.font.size = Pt(9)
        r1.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    doc.add_paragraph('')


def add_rank_boxes(doc, n=5):
    """Add ranking boxes (Deel 3)."""
    for i in range(1, n + 1):
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = 'Table Grid'
        c0 = tbl.rows[0].cells[0]
        c1 = tbl.rows[0].cells[1]
        c0.width = Cm(3.5)
        c1.width = Cm(12.5)

        p0 = c0.paragraphs[0]
        r0 = p0.add_run(f'Prioriteit {i}')
        r0.bold = True
        r0.font.size = Pt(9)
        r0.font.color.rgb = C_DARK_BLUE
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p1 = c1.paragraphs[0]
        r1 = p1.add_run('_' * 80)
        r1.font.size = Pt(9)
        r1.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    doc.add_paragraph('')


# ── LaTeX section extractors ──────────────────────────────────────────────────
def extract_between(text, start_pat, end_pat, flags=re.DOTALL):
    """Extract text between two regex patterns."""
    m = re.search(start_pat + r'(.*?)' + end_pat, text, flags)
    return m.group(1).strip() if m else ''


def extract_complaintbox(text, casus_n):
    r"""Extract the complaint text from \begin{complaintbox}...\end{complaintbox}."""
    # Find the N-th complaintbox
    boxes = re.findall(
        r'\\begin\{complaintbox\}.*?\\end\{complaintbox\}', text, re.DOTALL)
    if casus_n <= len(boxes):
        b = boxes[casus_n - 1]
        inner = re.sub(r'\\begin\{complaintbox\}(\[.*?\])?', '', b, flags=re.DOTALL)
        inner = re.sub(r'\\end\{complaintbox\}', '', inner)
        return latex2txt(inner)
    return ''


def extract_contextbox(text, casus_n):
    """Extract the context box (standardized symptoms/treatment goals)."""
    boxes = re.findall(
        r'\\begin\{contextbox\}.*?\\end\{contextbox\}', text, re.DOTALL)
    if casus_n <= len(boxes):
        b = boxes[casus_n - 1]
        inner = re.sub(r'\\begin\{contextbox\}', '', b)
        inner = re.sub(r'\\end\{contextbox\}', '', inner)
        return latex2txt(inner)
    return ''


def extract_monitorbox(text, casus_n):
    """Extract monitoring summary text."""
    boxes = re.findall(
        r'\\begin\{monitorbox\}.*?\\end\{monitorbox\}', text, re.DOTALL)
    if casus_n <= len(boxes):
        b = boxes[casus_n - 1]
        inner = re.sub(r'\\begin\{monitorbox\}', '', b)
        inner = re.sub(r'\\end\{monitorbox\}', '', inner)
        return latex2txt(inner)
    return ''


def extract_tikz_network(text, casus_n):
    """Extract network edge information from tikzpicture in Deel 3."""
    # Find tikzpictures in the Deel 3 section (skip the example one)
    tikz_blocks = re.findall(
        r'\\begin\{tikzpicture\}.*?\\end\{tikzpicture\}', text, re.DOTALL)
    # Index 0 is the example, casus_n+0 is the actual casus network
    # For casus_n, the actual casus tikz block is at index casus_n (0-indexed: example=0, casus1=1, casus2=2)
    # But Deel 3 section starts with example tikz, then casus 1 tikz, then casus 2 tikz
    # Actually tikzpictures appear in: Deel 3 example, Deel 3 casus 1, Deel 3 casus 2
    # There's one tikz in the main example (shared) and then the actual networks
    if casus_n + 0 < len(tikz_blocks):
        block = tikz_blocks[casus_n]  # 0=example, 1=casus1, 2=casus2
        # Extract BO and S node labels
        bo_nodes = re.findall(r'\\node\[(?:sm)?prnode\][^;]*\{BO-(\d+)\\\\([^}]+)\}', block)
        s_nodes = re.findall(r'\\node\[(?:sm)?crnode\][^;]*\{S-(\d+)\\\\([^}]+)\}', block)
        # Extract edges
        edges = re.findall(
            r'\\draw\[line width=([\d.]+)pt[^]]*draw=(PrimaryBlue|StrongRed)[^]]*\]\s*\(([^)]+)\).*?\(([^)]+)\)',
            block)
        result = []
        bo_dict = {n: lbl.replace('\\\\', ' ').strip() for n, lbl in bo_nodes}
        s_dict  = {n: lbl.replace('\\\\', ' ').strip() for n, lbl in s_nodes}
        # Describe edges as text
        for lw, color, src, dst in edges:
            direction = '→ positief verband' if 'PrimaryBlue' in color else '→ negatief verband'
            lw_f = float(lw)
            strength = 'zwak' if lw_f < 2.0 else ('matig' if lw_f < 3.5 else 'sterk')
            result.append(f'  • {src} {direction} {dst} ({strength})')
        return '\n'.join(result)
    return '(Zie PDF voor bipartiet netwerk)'


def extract_checkitems(text, casus_n):
    """Extract the 20 EMA check items for Deel 4."""
    # Find responsebox sections for Deel 4
    # Each responsebox in Deel 4 has \checkitemBFS entries
    # Find all blocks of \checkitemBFS
    check_blocks = list(re.finditer(
        r'(?=\\checkitemBFS\{1\})', text))
    if casus_n - 1 < len(check_blocks):
        start_pos = check_blocks[casus_n - 1].start()
        # Extract all checkitemBFS from this position
        chunk = text[start_pos:start_pos + 5000]
        items = re.findall(r'\\checkitemBFS\{(\d+)\}\{([^}]+)\}', chunk)
        return items[:20]
    return []


def extract_deel5_context(text, casus_n):
    """Extract the Deel 5 context table rows."""
    # Find tabular environments in Deel 5 contextbox sections
    ctx_blocks = list(re.finditer(
        r'\\begin\{contextbox\}.*?\\end\{contextbox\}', text, re.DOTALL))
    # Deel 5 contextboxes come after Deel 4 contextboxes
    # Deel 4 has casus_n contextboxes, Deel 5 starts at index (casus_n)
    # Actually: contextboxes appear in: Deel 2 (casus1, casus2), Deel 4 (casus1, casus2), Deel 5 (casus1, casus2)
    # That's 6 total. Deel 5 starts at index 4 (0-indexed).
    deel5_start = 4  # 0-indexed: indices 4 and 5 are Deel 5
    target_idx = deel5_start + (casus_n - 1)
    if target_idx < len(ctx_blocks):
        b = ctx_blocks[target_idx].group(0)
        # Extract tabular rows
        rows = re.findall(r'([^\\&\n][^\\\n]+?)\s*&\s*([^\\\n]+?)(?:\\\\|\\end)', b)
        result = []
        for label, value in rows:
            result.append((latex2txt(label.strip()), latex2txt(value.strip())))
        return result
    return []


# ── common instruction content ────────────────────────────────────────────────
INSTR_SECTION_TEXT = """\
PHOENIX (Personalized Hierarchical Optimization Engine for Navigating Insightful eXplorations) is een multi-agentsysteem ontwikkeld als onderdeel van een masterproef in de psychologie aan de Universiteit Gent. Het systeem verwerkt een vrije klachttekst en doorloopt vijf opeenvolgende klinische redeneerstappen die de kern vormen van een gepersonaliseerde, longitudinale digitale interventie.

Het theoretische kader van PHOENIX steunt op het netwerk-analytisch model van psychopathologie: psychische klachten worden niet als uitingen van een latente stoornis beschouwd, maar als een dynamisch netwerk van onderling samenhangende klachtcomponenten. Een sleutelbegrip daarbinnen is het onderscheid tussen symptomen (klacht- en toestandsdimensies die beschrijven wat er mis gaat bij de persoon) en behandelingsopties (modificeerbare gedragsvariabelen die de symptomen causaal beïnvloeden en als behandeldoel kunnen dienen)."""

KERNDEF_BOX = [
    "Symptoom (S): een klinisch relevante, momentaan aanwezige klacht- of toestandsdimensie van de persoon die herhaald meetbaar is via dagelijkse EMA. Symptomen zijn de uitkomstvariabelen in het netwerk -- ze beschrijven wat er mis gaat (bijv. inslaapproblemen, paniekepisoden, sombere stemming). Symptomen zijn geen gedragingen of interventies.",
    "",
    "Behandelingsoptie (BO): een modificeerbare gedrags- of procesvariabele die plausibel causaal ingrijpt op een of meerdere symptomen en dagelijks meetbaar is via een korte EMA-vraag. Behandelingsopties zijn de ingreepvariabelen in het netwerk -- ze beschrijven wat veranderbaar is (bijv. avondschermtijd, geplande exposure, aerobe beweging). Een behandelingsoptie is nooit een symptoom zelf, maar een gedrag, strategie of omgevingsfactor die het symptoom beïnvloedt.",
    "",
    "Voorbeelden ter verduidelijking:",
    "• Symptoom: inslaapproblemen, anticipatieangst, anergie, piekeren voor het slapengaan.",
    "• Behandelingsoptie: avondschermtijd (min), geplande exposurestap (ja/nee), bewegingsfrequentie (min), preslaap-offloading (ja/nee).",
]

EMA_BOX = [
    "EMA is een methode waarbij personen vaak meerdere keren per dag hun actuele toestand of gedrag rapporteren via een mobiele applicatie. In PHOENIX wordt EMA gebruikt om zowel symptomen als behandelingsopties dagelijks te monitoren over een periode van meerdere weken.",
    "",
    "Een goede EMA-variabele voldoet aan vier vereisten:",
    "• Dagelijks rapporteerbaar: via een korte vraag op de smartphone, bijv. (ja/nee, aantal, minuten of 0–10).",
    "• Dynamisch en veranderbaar: geen vaste trek, diagnose of stabiel achtergrondkenmerk.",
    "• Klinisch relevant: toont binnen-persoonsvariatie die therapeutisch informatief is.",
    "• Onderscheid symptoom vs. behandelingsoptie: klachtdimensies zijn symptomen; modificeerbare gedragingen zijn behandelingsopties.",
]

NETWERK_BOX = [
    "Na 21 dagen EMA-monitoring construeert PHOENIX een bipartiet netwerk: een graaf met twee kolommensets en gewogen verbindingen daartussen.",
    "• Linkse kolom – Behandelingsopties (BO): de modificeerbare variabelen.",
    "• Rechtse kolom – Symptomen (S): de klacht- en toestandsdimensies.",
    "• Kanten: de richting loopt van behandelingsoptie naar symptoom (BO → S). Een blauwe rand duidt op een positief verband (behandelingsoptie vergroot het symptoom), een rode rand op een negatief verband (behandelingsoptie verkleint het symptoom). Lijndikte is proportioneel aan de sterkte van het empirische verband.",
    "In Deel 3 rangschikt u de behandelingsopties op behandelprioriteit op basis van dit netwerk en de monitoringsamenvatting.",
]

HAPA_BOX = [
    "PHOENIX gebruikt het Health Action Process Approach (HAPA; Schwarzer, 1992) om de boodschap af te stemmen op de motivationele fase van de persoon:",
    "• Pre-intentionele fase: de persoon is nog niet gemotiveerd om te veranderen → focus op risicobewustzijn en uitkomstverwachting.",
    "• Intentionele fase: de persoon wil veranderen maar heeft nog geen concreet plan → focus op doelstelling en actieplanning.",
    "• Actie-/onderhoudsfase: de persoon probeert al te veranderen → focus op copingplanning en zelfeffectiviteitsondersteuning.",
    "",
    "De boodschap voldoet aan:",
    "• Lengte: 2–4 zinnen, compact genoeg voor een mobiel scherm.",
    "• Toon: warm, direct, professioneel – geen klinisch jargon of diagnostische labels.",
    "• Inhoud: adresseert het primaire behandeldoel en de voornaamste barrière; bevat een concrete, eerstvolgende actie.",
    "• Perspectief: tweede persoon ('jij' of formeel 'u').",
]


def get_werkwijze_text(is_two_case: bool) -> list:
    case_phrase = "voor beide casussen" if is_two_case else ""
    step1 = f"Werk sequentieel: Deel 1 → Deel 5. Ga pas naar een volgend deel wanneer het huidige deel volledig is afgewerkt{' ' + case_phrase if case_phrase else ''}."
    return [
        f"1. {step1}",
        "2. Gebruik geen generatieve AI, schrijfhulpmiddelen, richtlijnen of collegaoverleg. Extern gebruik ondermijnt de methodologische validiteit en de blind scoringswaarde van de studie.",
        "3. Gebruik in latere delen uitsluitend de meegeleverde gestandaardiseerde context. Die is bewust vastgezet zodat alle deelnemers op identieke input reageren en vergelijkbaar zijn.",
        "4. Herwerk eerdere antwoorden niet retroactief nadat u latere context hebt gezien.",
        "5. Noteer of typ rechtstreeks in de voorziene antwoordzones. Onleesbare of ambigu geformuleerde antwoorden bemoeilijken latere blinde beoordeling.",
        "6. Werk ook bij beperkte informatie. Klinische teksten bevatten vaak onvoldoende informatie voor definitieve conclusies. Geef altijd een antwoord op basis van de beschikbare informatie en uw klinisch oordeel. Laat geen velden blanco zonder reden.",
    ]


# ── Deel instruction texts ────────────────────────────────────────────────────
DEEL1_INSTR = [
    "Opdracht: identificeer de belangrijkste actuele klacht- en toestandsdimensies (symptomen) in de klachttekst en noteer voor elke dimensie uitsluitend een kort symptoomlabel.",
    "",
    "Wat is een symptoom? Een symptoom is een klacht- of toestandsdimensie die beschrijft wat er mis gaat bij de persoon. Het is geen behandeling, geen gedrag en geen oorzaak, maar een actuele toestandsbeschrijving:",
    "• momenteel aanwezig bij de persoon (niet hypothetisch of anamnestisch),",
    "• onderscheidbaar van andere klachtdimensies (niet overlappend),",
    "• in principe herhaald meetbaar via dagelijkse zelfrapportage (EMA),",
    "• op DSM-5-TR/ICD-11-niveau: concreet genoeg om afzonderlijk te meten (bijv. 'inslaapmoeilijkheden', niet het brede 'slaapstoornis'), maar niet te atomair (niet 'wakker om 3u00' als apart symptoom).",
    "",
    "Antwoordformat: label van 2–5 woorden, zonder beschrijvende zin. Noteer per casus 2–6 symptomen. Laat ongebruikte velden leeg.",
]

DEEL2_INSTR = [
    "Opdracht: genereer 3–5 behandelingsoptielabels die samen het initieel observatiemodel voor deze casus vormen.",
    "",
    "Wat is een behandelingsoptie in deze context? Een behandelingsoptie is een modificeerbare gedrags- of procesvariabele die beschrijft wat de persoon kan veranderen:",
    "• Modificeerbaar: de persoon (of therapeut) kan er rechtstreeks op ingrijpen.",
    "• EMA-geschikt: dagelijks meetbaar via een eenvoudige vraag op de smartphone (ja/nee, aantal, minuten of 0–10 schaal).",
    "• Causaal plausibel: klinisch aannemelijk dat de behandelingsoptie de symptomen beïnvloedt.",
    "• Geen symptoom: klachtniveaus of diagnostische trekken zijn geen behandelingsopties.",
    "",
    "Antwoordformat: enkel een label van 2–6 woorden, zonder meetdefinitie of toelichting. Laat ongebruikte velden leeg.",
]

DEEL3_INSTR = [
    "Opdracht: rangschik de 5 gestandaardiseerde behandelingsopties van hoogste naar laagste behandelprioriteit.",
    "",
    "Een behandeldoel is de behandelingsoptie die, als zij veranderd wordt, naar verwachting de sterkste vermindering van de symptomen oplevert. Prioriteer op basis van monitoringbewijs (frequentie en ernst van het gedragspatroon), klinische modificeerbaarheid (haalbaarheid voor deze persoon) en netwerkimpact (invloed op meerdere symptomen).",
    "",
    "Bipartiet netwerk lezen:",
    "• Behandelingsopties (groen/links); symptomen (blauw/rechts).",
    "• Blauwe rand: positief verband (meer BO → meer S).",
    "• Rode rand: negatief verband (meer BO → minder S).",
    "• Lijndikte: proportioneel aan de sterkte van het verband (21-daagse EMA-data).",
    "",
    "Antwoordformat: vul alle 5 prioriteitslijnen in, van rangorde 1 (hoogste) tot 5 (laagste).",
]

DEEL4_INSTR = [
    "Opdracht: selecteer per casus exact 6 EMA-items: 2 sub-behandelingsopties per behandeldoel (3 behandeldoelen × 2 = 6 items).",
    "",
    "Sub-behandelingsopties zijn de concrete gedragsoperationalisering van een abstract behandeldoel (bijv. 'Slaapkwaliteit' → 'schermvrij interval voor slapengaan (min)').",
    "",
    "• Selecteer per behandeldoel de 2 items die het domein het meest direct en precies meten.",
    "• Vermijd items die het domein slechts zijdelings raken.",
    "• Alle 20 items zijn behandelingsoptie-type EMA-items (modificeerbare gedragingen – geen symptomen).",
    "",
    "Antwoordformat: vink exact 6 items aan. Geen toelichting vereist.",
    "In het Word-document kunt u de vakjes (☐) omcirkelen of markeren in plaats van aanvinken.",
]

DEEL5_INSTR = [
    "Opdracht: schrijf een korte, rechtstreeks tot de persoon gerichte coachingsboodschap die in de mobiele applicatie verschijnt.",
    "",
    "Gebruik het HAPA-kader (zie kader hieronder) om de toon en inhoud van uw boodschap te sturen. U hoeft de fase niet expliciet te benoemen.",
]


# ── main generation function ──────────────────────────────────────────────────
def extract_metadata(text: str):
    """Extract key metadata fields from the LaTeX source."""
    hcp_code = re.search(r'\\textbf\{(HCP-PRE-\d+)\}', text)
    hcp_code = hcp_code.group(1) if hcp_code else 'HCP-PRE-??'

    # Case assignments from the deelnemerscode box
    cases = re.findall(r'\\textbf\{\\color\{PrimaryBlue\}(C\d+)\}\s*\(([^)]+)\)', text)

    # Geschatte duur
    duur = re.search(r'Geschatte duur & ([^\\]+)\\\\', text)
    duur = duur.group(1).strip() if duur else ''

    # Email subject
    subj = re.search(r'met onderwerp:\\quad \\texttt\{([^}]+)\}', text)
    subj = subj.group(1) if subj else ''

    return hcp_code, cases, duur, subj


def extract_casus_labels(text, casus_n, num_casussen):
    """
    Extract all case-specific content blocks for a given casus number.
    Returns dict with keys: subsection_title, complaint, d2_summary, d2_context,
    d3_monitor, d3_bos, d4_goals, d4_items, d5_context
    """
    # Find subsection headers for this casus
    pattern = rf'\\subsection\*\{{Casus {casus_n}: ([^}}]+?)\}}'
    headers = re.findall(pattern, text)

    return headers  # For now just return the headers


def generate_word_doc(tex_path: str, out_path: str, is_two_case: bool) -> None:
    with open(tex_path, encoding='utf-8') as f:
        text = f.read()

    hcp_code, cases, duur, email_subj = extract_metadata(text)

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Page margins
    section = doc.sections[0]
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)
    section.top_margin    = Cm(2.4)
    section.bottom_margin = Cm(2.4)

    # ── TITLE PAGE ──────────────────────────────────────────────────────────
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run('PHOENIX evaluatiestudie')
    tr.font.size = Pt(22)
    tr.bold = True
    tr.font.color.rgb = C_DARK_BLUE

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2r = t2.add_run('Fase 1 – onafhankelijke expertgeneratie')
    t2r.font.size = Pt(15)
    t2r.font.color.rgb = C_PRIMARY

    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t3r = t3.add_run('Instrument voor Zorgprofessionals')
    t3r.italic = True
    t3r.font.size = Pt(12)

    doc.add_paragraph('')

    # Deelnemerscode box
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'DBEAFE')
    tcPr.append(shd)

    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cp.add_run(f'Deelnemerscode:  {hcp_code}')
    cr.bold = True
    cr.font.size = Pt(14)
    cr.font.color.rgb = C_DARK_BLUE

    # Cases
    if is_two_case and len(cases) >= 2:
        cp2 = cell.add_paragraph()
        cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp2r = cp2.add_run(
            f'Toegewezen casussen:  {cases[0][0]} ({cases[0][1]})  en  {cases[1][0]} ({cases[1][1]})')
        cp2r.font.size = Pt(11)
        cp2r.font.color.rgb = C_SLATE
    elif cases:
        cp2 = cell.add_paragraph()
        cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp2r = cp2.add_run(f'Toegewezen casus:  {cases[0][0]} ({cases[0][1]})')
        cp2r.font.size = Pt(11)
        cp2r.font.color.rgb = C_SLATE

    doc.add_paragraph('')

    # Info table
    info_tbl = doc.add_table(rows=6, cols=2)
    info_tbl.style = 'Table Grid'
    info = [
        ('Studie',       'Evaluatie van de klinische kwaliteit van een ontologiegebaseerd multi-agentsysteem voor gepersonaliseerde digitale geestelijke gezondheidszorg'),
        ('Instelling',   'Universiteit Gent – Faculteit Psychologie en Pedagogische Wetenschappen'),
        ('Onderzoeker',  'Stijn Van Severen (masterproefstudent)'),
        ('Promotoren',   'Prof. Dr. Geert Crombez; Dr. Annick De Paepe'),
        ('Contact',      'stijn.vanseveren@ugent.be'),
        ('Geschatte duur', duur),
    ]
    for i, (label, value) in enumerate(info):
        row = info_tbl.rows[i]
        lc = row.cells[0]
        vc = row.cells[1]
        lr = lc.paragraphs[0].add_run(label)
        lr.bold = True
        lr.font.size = Pt(10)
        vr = vc.paragraphs[0].add_run(value)
        vr.font.size = Pt(10)

    doc.add_paragraph('')

    # Doel box
    # Extract doel text from LaTeX
    doel_raw = extract_between(text,
        r'\\textbf\{\\color\{PrimaryBlue\}Doel van deze bundel\}',
        r'\\end\{tcolorbox\}\s*\\end\{center\}\s*\\vspace')
    doel_txt = latex2txt(doel_raw)
    add_box(doc, 'Doel van deze bundel', doel_txt.split('\n'), C_PRIMARY, 'F8FAFC')

    # Vertrouwelijkheid
    conf = "Uw antwoorden worden voor analyse geanonimiseerd en uitsluitend gebruikt binnen deze masterproefstudie. Deelname is vrijwillig; u kan zich op elk moment terugtrekken door contact op te nemen met de onderzoeker."
    add_box(doc, 'Vertrouwelijkheid', [conf], C_AMBER, 'FEF3C7')

    doc.add_page_break()

    # ── INSTRUCTION PAGE ────────────────────────────────────────────────────
    add_heading(doc, '1  Instructiepagina', level=1)
    add_heading(doc, 'Het systeem en de context van deze studie', level=2, color=C_PRIMARY)
    add_para(doc, INSTR_SECTION_TEXT, size=10)
    doc.add_paragraph('')

    add_box(doc, 'Kerndefinities: symptoom versus behandelingsoptie', KERNDEF_BOX, C_SLATE)
    add_box(doc, 'Ecological Momentary Assessment (EMA) – basisprincipes', EMA_BOX, C_SLATE)
    add_box(doc, 'Het bipartiet netwerk: structuur van het observatiemodel', NETWERK_BOX, C_SLATE)

    # Time table
    add_heading(doc, 'Overzicht van de vijf taken', level=2, color=C_PRIMARY)
    time_tbl = doc.add_table(rows=8 if is_two_case else 8, cols=4)
    time_tbl.style = 'Table Grid'
    headers_t = ['Stap', 'Klinische taak', 'Wat u doet', 'Richttijd']
    rows_t = [
        ('--',  'Instructies lezen',       'Lees de instructiepagina aandachtig door',                            '≈ 5 min'),
        ('1',   'Operationalisering',       'Identificeer 2–6 symptoomlabels (klachtdimensies)',                   '≈ 6 min' if is_two_case else '≈ 3 min'),
        ('2',   'Initieel observatiemodel', 'Genereer 3–5 behandelingsoptielabels (modificeerbaar, EMA-geschikt)', '≈ 6 min' if is_two_case else '≈ 3 min'),
        ('3',   'Behandeldoelprioritering', 'Rangschik alle 5 behandelingsopties van hoog naar laag',              '≈ 7 min' if is_two_case else '≈ 4 min'),
        ('4',   'Verfijnd observatiemodel', 'Selecteer exact 6 EMA-items (2 per behandeldoel) uit de lijst',      '≈ 8 min' if is_two_case else '≈ 4 min'),
        ('5',   'Mobiele coaching',         'Schrijf een korte patiëntgerichte boodschap voor de app',            '≈ 8 min' if is_two_case else '≈ 4 min'),
        ('',    'Totaal (' + ('2 casussen' if is_two_case else '1 casus') + ')', '',  '40–50 min' if is_two_case else '20–25 min'),
    ]
    for j, h in enumerate(headers_t):
        c = time_tbl.rows[0].cells[j]
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9)
    for i, row_data in enumerate(rows_t):
        row = time_tbl.rows[i + 1]
        for j, val in enumerate(row_data):
            c = row.cells[j]
            r = c.paragraphs[0].add_run(val)
            r.font.size = Pt(9)
            if i == len(rows_t) - 1:
                r.bold = True

    doc.add_paragraph('')
    add_box(doc, 'Werkwijze – strikt te volgen',
            get_werkwijze_text(is_two_case), C_SLATE)

    doc.add_page_break()

    # ── DETERMINE HOW MANY CASUSSEN ─────────────────────────────────────────
    num_casussen = 2 if is_two_case else 1
    casus_list = list(range(1, num_casussen + 1))

    # ── DEEL 1 ──────────────────────────────────────────────────────────────
    add_heading(doc, '2  Deel 1: Operationalisering van de mentale gezondheidstoestand', level=1)
    add_box(doc, 'Instructies Deel 1', DEEL1_INSTR, C_PRIMARY, 'EFF6FF')

    # Example box
    add_box(doc, 'Uitgewerkt voorbeeld – Deel 1 – Nadia (38-jarige verpleegkundige)  |  Hoe formuleert u een correct symptoomlabel?',
        [
            "Klacht (verkorte versie): 'De afgelopen vijf maanden slaap ik slecht: ik lig lang wakker en word vroeg wakker. Overdag ben ik moe en prikkelbaar. Ik beweeg bijna niet meer en zie vrienden nauwelijks nog.'",
            "",
            "Correct ingevulde symptoomlabels:",
            "  S-1  Inslaapmoeilijkheden    ('lig lang wakker' → afzonderlijk meetbaar)",
            "  S-2  Vroegochtendontwaak    ('word vroeg wakker' → DSM-5-TR: early morning awakening; onderscheidbaar van S-1)",
            "  S-3  Vermoeidheid overdag    ('moe overdag' → toestandsdimensie)",
            "  S-4  Prikkelbaarheid         ('prikkelbaar' → emotioneel symptoom, los van S-3)",
            "  S-5  Sociale terugtrekking   ('zie vrienden nauwelijks' → klachtpatroon)",
            "",
            "Dit zijn GEEN symptoomlabels (horen in Deel 2 als behandelingsoptie):",
            "  'Weinig bewegen' → modificeerbaar gedrag, geen klacht; hoort bij Deel 2.",
            "  'Slaapproblemen' → te breed: splits naar meetbare dimensies (zie S-1 en S-2).",
        ], C_PURPLE, 'EDE9FE')

    # Per casus
    for cn in casus_list:
        complaint = extract_complaintbox(text, cn)
        add_heading(doc, f'Casus {cn}', level=2, color=C_PRIMARY)
        if complaint:
            add_box(doc, f'Casusvignet – Casus {cn}', [complaint], C_PRIMARY, 'EFF6FF')

        # Answer section
        add_answer_box_with_slots(doc, 'Uw antwoord – Deel 1',
            'Noteer 2–6 symptoomlabels voor deze casus. Gebruik enkel korte labels (2–5 woorden); voeg geen beschrijving toe.',
            slot_label='Symptoom', n_slots=6)

    doc.add_page_break()

    # ── DEEL 2 ──────────────────────────────────────────────────────────────
    add_heading(doc, '3  Deel 2: Initieel observatiemodel (bipartiet netwerk)', level=1)
    add_box(doc, 'Instructies Deel 2', DEEL2_INSTR, C_FOREST, 'ECFDF5')

    add_box(doc, 'Uitgewerkt voorbeeld – Deel 2 – Nadia (38-jarige verpleegkundige)  |  Hoe genereert u behandelingsoptielabels?',
        [
            "Gestandaardiseerde symptomen uit Deel 1 (voorbeeld):",
            "S-1: Inslaapmoeilijkheden;  S-2: Vroegochtendontwaak;  S-3: Vermoeidheid overdag;  S-4: Prikkelbaarheid;  S-5: Sociale terugtrekking.",
            "",
            "Correct ingevulde behandelingsoptielabels:",
            "  BO-1  Avondschermgebruik       (modificeerbaar gedrag; EMA: minuten voor slapengaan)",
            "  BO-2  Lichaamsbeweging          (modificeerbaar; EMA: minuten actief bewegen per dag)",
            "  BO-3  Sociaal contact            (modificeerbaar; EMA: ja/nee contact gezocht op eigen initiatief)",
            "",
            "Dit zijn GEEN behandelingsopties (zijn symptomen, horen in Deel 1):",
            "  'Vermoeidheid' → klachtdimensie (= S-2), geen behandeldoel.",
            "  'Slechte slaap' → klachtdimensie (= S-1), geen behandeldoel.",
        ], C_PURPLE, 'EDE9FE')

    for cn in casus_list:
        complaint = extract_complaintbox(text, cn)
        context = extract_contextbox(text, cn)
        add_heading(doc, f'Casus {cn} – Deel 2', level=2, color=C_PRIMARY)
        if complaint:
            add_box(doc, f'Verkorte klachtomschrijving – Casus {cn}', [complaint.split('\n')[0]], C_PRIMARY, 'EFF6FF')
        if context:
            add_box(doc, 'Gestandaardiseerde symptomen uit Deel 1', context.split('\n'), C_FOREST, 'ECFDF5')
        add_answer_box_with_slots(doc, 'Uw antwoord – Deel 2',
            'Noteer 3–5 behandelingsoptielabels voor een initieel observatiemodel. Zorg dat elk label later dagelijks via een mobiele app meetbaar kan worden gemaakt.',
            slot_label='Behandelingsoptie', n_slots=5)

    doc.add_page_break()

    # ── DEEL 3 ──────────────────────────────────────────────────────────────
    add_heading(doc, '4  Deel 3: Behandeldoelprioritering via het bipartiet netwerk', level=1)
    add_box(doc, 'Instructies Deel 3', DEEL3_INSTR, C_AMBER, 'FEF3C7')

    add_box(doc, 'Uitgewerkt voorbeeld – Deel 3 – Nadia (38-jarige verpleegkundige)  |  Hoe leest u het netwerk en rangschikt u de behandelingsopties?',
        [
            "Monitoring (voorbeeld): Schermtijd voor slapengaan: 18/21 avonden actief (gem. 65 min). Lichaamsbeweging: 0,8 sessies/week. Sociaal contact op eigen initiatief: 0,9/week.",
            "",
            "Bipartiet netwerk (voorbeeld – blauw = positief, rood = negatief):",
            "  BO-1 (Avondschermgebruik) → S-1 (Inslaap) [blauw, sterk]; → S-2 (Vermoeidheid) [blauw, matig]",
            "  BO-2 (Lichaamsbeweging)   → S-1 [rood, matig]; → S-2 [rood, zwak]",
            "  BO-3 (Sociaal contact)    → S-3 (Sociale terugtrekking) [rood, sterk]",
            "",
            "Prioriteringsredenering:",
            "  1. BO-1 (Avondschermgebruik): raakt S-1 en S-2; monitoring 18/21 avonden → prioriteit 1.",
            "  2. BO-2 (Lichaamsbeweging): beschermend voor S-1 en S-2; frequentie extreem laag → prioriteit 2.",
            "  3. BO-3 (Sociaal contact): sterk negatief verband met S-3; laag (0,9/week) → prioriteit 3.",
        ], C_PURPLE, 'EDE9FE')

    # Get Deel 3-specific contextboxes (these are indices 0..3 for Deel 2, so we
    # need the actual Casus-specific monitoring and network data)
    for cn in casus_list:
        add_heading(doc, f'Casus {cn} – Deel 3', level=2, color=C_PRIMARY)

        # Extract complaint summary (second set of complaintboxes in Deel 3)
        # Deel 3 complaintboxes are the 3rd and 4th occurrences
        d3_complaint_idx = 2 + cn - 1  # 0-indexed
        all_complaints = re.findall(
            r'\\begin\{complaintbox\}.*?\\end\{complaintbox\}', text, re.DOTALL)
        if d3_complaint_idx < len(all_complaints):
            d3_c = latex2txt(re.sub(r'\\begin\{complaintbox\}(\[.*?\])?|\\end\{complaintbox\}', '',
                                    all_complaints[d3_complaint_idx], flags=re.DOTALL))
            if d3_c:
                add_box(doc, f'Casus {cn}', [d3_c], C_PRIMARY, 'EFF6FF')

        # Monitoring summary
        monitor = extract_monitorbox(text, cn)
        if monitor:
            add_box(doc, '21-daagse monitoringsamenvatting', [monitor], C_SLATE, 'F8FAFC')

        # Network: extract BO labels from tikzpicture
        # Find the casus-specific tikzpictures (skip example = index 0)
        tikz_idx = cn  # 0=example, 1=casus1, 2=casus2
        all_tikz = re.findall(
            r'\\begin\{tikzpicture\}.*?\\end\{tikzpicture\}', text, re.DOTALL)
        if tikz_idx < len(all_tikz):
            block = all_tikz[tikz_idx]
            bo_labels = re.findall(r'\\node\[prnode\][^;]*\{BO-(\d+)\\\\([^}]+)\}', block)
            s_labels  = re.findall(r'\\node\[crnode\][^;]*\{S-(\d+)\\\\([^}]+)\}', block)
            edges     = re.findall(
                r'\\draw\[line width=([\d.]+)pt[^\]]*draw=(PrimaryBlue|StrongRed)[^\]]*\]\s*\((\w+)\.east\).*?\((\w+)\.west\)',
                block)

            # Build BO/S label maps
            bo_map = {f'p{n}': lbl.replace('\\\\', ' ') for n, lbl in bo_labels}
            s_map  = {f'cr{n}': lbl.replace('\\\\', ' ') for n, lbl in s_labels}

            network_lines = ['Bipartiet netwerk (blauw = positief verband; rood = negatief verband):']
            network_lines.append(
                'Behandelingsopties (links):  ' +
                ', '.join(f'BO-{n}: {lbl}' for n, lbl in bo_labels))
            network_lines.append(
                'Symptomen (rechts):  ' +
                ', '.join(f'S-{n}: {lbl}' for n, lbl in s_labels))
            network_lines.append('')
            for lw, color, src, dst in edges:
                src_lbl = bo_map.get(src, src)
                dst_lbl = s_map.get(dst, dst)
                kind = 'positief' if 'Blue' in color else 'negatief'
                lw_f = float(lw)
                strength = '(zwak)' if lw_f < 2.0 else ('(matig)' if lw_f < 3.5 else '(sterk)')
                network_lines.append(f'  • {src_lbl} → {dst_lbl}  [{kind}] {strength}')

            add_box(doc, 'Bipartiet netwerk', network_lines, C_SLATE, 'F0F4FF')

        # Response area
        add_heading(doc, 'Uw antwoord – Deel 3', level=3, color=C_PURPLE)

        # Extract BO list from response box
        bos_raw = re.findall(r'\\textbf\{BO-\d+: [^}]+\}', text)
        # Take the casus-specific set
        # Deel 3 Casus N has 5 BOs listed in the responsebox
        # Extract from around the N-th set of rankboxes
        add_para(doc, 'Rangschik alle 5 behandelingsopties van hoogste naar laagste behandelprioriteit:', size=10)
        add_rank_boxes(doc, 5)

    doc.add_page_break()

    # ── DEEL 4 ──────────────────────────────────────────────────────────────
    add_heading(doc, '5  Deel 4: Verfijnd observatiemodel – selectie van sub-behandelingsopties', level=1)
    add_box(doc, 'Instructies Deel 4', DEEL4_INSTR, C_PURPLE, 'F3E8FF')

    add_box(doc, 'Uitgewerkt voorbeeld – Deel 4 – Nadia (38-jarige verpleegkundige)  |  Hoe selecteert u de meest passende EMA-items per behandeldoel?',
        [
            "Gestandaardiseerde behandeldoelen (abstract, voorbeeld):",
            "  1. Slaapkwaliteit",
            "  2. Lichamelijke activiteit",
            "  3. Sociaal contact",
            "",
            "Selectielogica (doel 1 = Slaapkwaliteit):",
            "  ✓ 'Schermvrij interval direct voor slapengaan (min)' → juist (directe operationalisering)",
            "  ✓ 'Vaste bedtijd aangehouden (ja/nee)' → juist (complementaire meting: regelmaat)",
            "  × 'Cafeïne-inname na 15.00 uur (aantal)' → niet juist (raakt slaap slechts zijdelings)",
            "  × 'Maaltijd op vaste tijden (ja/nee)' → niet juist (hoort bij Voedingskwaliteit)",
            "",
            "Correct ingevuld (6 items totaal):",
            "  Doel 1 → 'Schermvrij interval voor slapengaan (min)' + 'Vaste bedtijd aangehouden (ja/nee)'",
            "  Doel 2 → 'Duur actieve beweging vandaag (min)' + 'Beweegactiviteit uitgevoerd (ja/nee)'",
            "  Doel 3 → 'Bewust sociaal contact gezocht vandaag (ja/nee)' + 'Sociale activiteit bijgewoond of gepland (ja/nee)'",
        ], C_PURPLE, 'EDE9FE')

    for cn in casus_list:
        add_heading(doc, f'Casus {cn} – Deel 4', level=2, color=C_PRIMARY)

        # Context box for Deel 4
        # Deel 4 contextboxes are indices 2 and 3 (after Deel 2 contextboxes at 0 and 1)
        d4_ctx_idx = 2 + cn - 1  # 0-indexed
        all_ctx = re.findall(
            r'\\begin\{contextbox\}.*?\\end\{contextbox\}', text, re.DOTALL)
        if d4_ctx_idx < len(all_ctx):
            d4_ctx = latex2txt(re.sub(
                r'\\begin\{contextbox\}|\\end\{contextbox\}', '',
                all_ctx[d4_ctx_idx], flags=re.DOTALL))
            if d4_ctx:
                add_box(doc, 'Gestandaardiseerde behandeldoelen (abstract)', d4_ctx.split('\n'), C_FOREST, 'ECFDF5')

        # Check items
        items = extract_checkitems(text, cn)
        add_para(doc, 'Opdracht: selecteer exact 6 EMA-items (2 per behandeldoel). Omcirkel of markeer de vakjes (☐) in dit Word-document, of typ een X bij de geselecteerde items.', size=10, italic=True)
        doc.add_paragraph('')

        if items:
            tbl = doc.add_table(rows=len(items) + 1, cols=2)
            tbl.style = 'Table Grid'
            # Header
            h0 = tbl.rows[0].cells[0].paragraphs[0].add_run('Nr.')
            h0.bold = True
            h0.font.size = Pt(9)
            h1 = tbl.rows[0].cells[1].paragraphs[0].add_run('EMA-item')
            h1.bold = True
            h1.font.size = Pt(9)
            for i, (num, item_txt) in enumerate(items):
                row = tbl.rows[i + 1]
                c0 = row.cells[0]
                c1 = row.cells[1]
                r0 = c0.paragraphs[0].add_run(f'{num}.  ☐')
                r0.font.size = Pt(9)
                r1 = c1.paragraphs[0].add_run(item_txt)
                r1.font.size = Pt(9)

        add_para(doc, '')
        add_para(doc, f'Totaal geselecteerd: ______ / 6', size=10, bold=True)
        doc.add_paragraph('')

    doc.add_page_break()

    # ── DEEL 5 ──────────────────────────────────────────────────────────────
    add_heading(doc, '6  Deel 5: Gepersonaliseerde mobiele coachingsboodschap (HAPA-kader)', level=1)
    add_box(doc, 'Instructies Deel 5', DEEL5_INSTR, C_FOREST, 'ECFDF5')
    add_box(doc, 'HAPA-kader', HAPA_BOX, C_SLATE, 'F8FAFC')

    add_box(doc, 'Uitgewerkt voorbeeld – Deel 5 – Nadia (38-jarige verpleegkundige)  |  Hoe formuleert u een effectieve coachingsboodschap?',
        [
            "Context (voorbeeld):",
            "  Primair behandeldoel: avondschermgebruik verminderen",
            "  Voornaamste barrière: gewoontekracht – schermgebruik is de enige manier om te ontspannen (lage zelfeffectiviteit)",
            "  HAPA-fase: intentioneel (wil veranderen, maar heeft nog geen concreet plan)",
            "",
            "Voorbeeldboodschap:",
            "  'Je weet dat je avondscherm je slaap in de weg staat, maar het voelt nog als de makkelijkste manier om te ontspannen na een drukke dag. Leg vanavond je telefoon om 21.30 uur in een andere kamer en vul die laatste twintig minuten met iets anders – een boek, muziek, of gewoon even niets. Eén avond is genoeg om te merken dat het kan.'",
            "",
            "Waarom werkt deze boodschap? De boodschap erkent de barrière (gewoontekracht), geeft een concrete actie met tijdstip (21.30u), verlaagt de drempel en vergroot de zelfeffectiviteit.",
        ], C_PURPLE, 'EDE9FE')

    for cn in casus_list:
        add_heading(doc, f'Casus {cn} – Deel 5', level=2, color=C_PRIMARY)

        # Context table for Deel 5 (indices 4 and 5 across all contextboxes)
        all_ctx = re.findall(
            r'\\begin\{contextbox\}.*?\\end\{contextbox\}', text, re.DOTALL)
        d5_ctx_idx = 4 + cn - 1  # 0-indexed
        if d5_ctx_idx < len(all_ctx):
            d5_block = all_ctx[d5_ctx_idx]
            # Extract tabular rows from contextbox
            rows_raw = re.findall(
                r'([A-Z][^\\\n&]{3,}) &\s*([^\\\n]+?)(?:\\\\|\n\\end)', d5_block)
            ctx_lines = []
            for label, value in rows_raw:
                label_c = latex2txt(label.strip())
                value_c = latex2txt(value.strip())
                if label_c and value_c:
                    ctx_lines.append(f'{label_c}:  {value_c}')
            if ctx_lines:
                add_box(doc, f'Context Casus {cn}', ctx_lines, C_FOREST, 'ECFDF5')

        add_answer_box(doc, 'Uw coachingsboodschap', lines=5)

    doc.add_page_break()

    # ── AFRONDING ────────────────────────────────────────────────────────────
    add_heading(doc, '7  Afronding en terugbezorging', level=1)
    if is_two_case:
        add_para(doc, 'Dank u voor het invullen van alle vijf delen voor uw twee toegewezen casussen (Casus 1 en Casus 2).')
    else:
        add_para(doc, 'Dank u voor het invullen van alle vijf delen voor uw toegewezen casus.')

    checklist_items = [
        'Ik heb in Deel 1 symptoomlabels ingevuld.',
        'Ik heb in Deel 2 behandelingsoptielabels ingevuld.',
        'Ik heb in Deel 3 alle 5 behandelingsopties gerangschikt.',
        'Ik heb in Deel 4 exact 6 EMA-items geselecteerd.',
        'Ik heb in Deel 5 een mobiele coachingsboodschap geschreven.',
        'Mijn antwoorden weerspiegelen mijn eigen klinische oordeel zonder gebruik van generatieve AI of andere externe hulp.',
        'Ik begrijp dat mijn antwoorden geanonimiseerd worden voor analyse.',
    ]
    add_box(doc, 'Checklist voor terugbezorging',
            [f'☐  {item}' for item in checklist_items], C_SLATE)

    add_para(doc, f'Bezorg het ingevulde document terug via e-mail aan: stijn.vanseveren@ugent.be')
    add_para(doc, f'Onderwerp: {email_subj}')

    doc.add_paragraph('')
    add_box(doc, '',
        ['Na ontvangst worden uw antwoorden geanonimiseerd en opgenomen in het expertreferentiecorpus voor de latere blind evaluatie. Hartelijk dank voor uw bijdrage aan dit onderzoek.'],
        C_PRIMARY, 'DBEAFE')

    # ── SAVE ────────────────────────────────────────────────────────────────
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    print(f'  ✓  Saved: {out_path}')


def add_answer_box_with_slots(doc, label, instruction, slot_label, n_slots):
    """Answer box with instruction + N labeled slots."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFFDE7')
    tcPr.append(shd)

    lp = cell.paragraphs[0]
    lr = lp.add_run(label)
    lr.bold = True
    lr.font.size = Pt(10)
    lr.font.color.rgb = C_PURPLE

    ip = cell.add_paragraph(instruction)
    ip.runs[0].font.size = Pt(9)
    ip.runs[0].italic = True

    for i in range(1, n_slots + 1):
        sp = cell.add_paragraph()
        sp.paragraph_format.space_before = Pt(4)
        sr = sp.add_run(f'{slot_label} {i}   Label: ')
        sr.bold = True
        sr.font.size = Pt(9)
        sr.font.color.rgb = C_SLATE
        sp.add_run('_' * 60).font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    doc.add_paragraph('')


# ── entry point ───────────────────────────────────────────────────────────────
if __name__ == '__main__':
    script_dir = os.path.dirname(os.path.abspath(__file__))

    # 2_cases_per_HCP
    src_2c = os.path.join(script_dir, '2_cases_per_HCP')
    for i in range(1, 6):
        tex = os.path.join(src_2c, f'HCP_{i}', 'main.tex')
        docx_out = os.path.join(src_2c, f'HCP_{i}', 'main.docx')
        print(f'\nGenerating 2-case Word doc for HCP_{i} …')
        generate_word_doc(tex, docx_out, is_two_case=True)

    # 1_case_per_HCP
    src_1c = os.path.join(script_dir, '1_case_per_HCP')
    for i in range(1, 11):
        code = f'{i:02d}'
        tex = os.path.join(src_1c, f'HCP_{code}', 'main.tex')
        docx_out = os.path.join(src_1c, f'HCP_{code}', 'main.docx')
        print(f'\nGenerating 1-case Word doc for HCP_{code} …')
        generate_word_doc(tex, docx_out, is_two_case=False)

    print('\nAll Word docs generated.')
