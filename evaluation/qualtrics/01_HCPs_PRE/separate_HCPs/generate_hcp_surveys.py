"""
Genereer de vijf afzonderlijke PRE-instrumenten voor zorgprofessionals.

Elke bundel bevat uitsluitend de twee aan die deelnemer toegewezen casussen en
is volledig geformuleerd in professioneel Nederlands. De documenten zijn
bedoeld als finale distributieversie voor dataverzameling.
"""

from __future__ import annotations

import subprocess
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


CASES: dict[str, dict] = {
    "C01": dict(
        label="C01",
        profile="28-jarige softwareontwikkelaar",
        duration="2 maanden",
        short_desc=(
            "Inslaapproblemen, cognitieve hyperarousal bij het slapengaan, "
            "sociale anhedonie en stressgerelateerde spierspanning."
        ),
        vignette=(
            r"``De afgelopen twee maanden geraak ik 's avonds heel moeilijk in slaap. "
            r"Mijn hoofd blijft maar doorgaan over alles wat ik nog moet doen -- deadlines, "
            r"onafgewerkte taken, gesprekken die ik nog moet voeren. Zelfs als ik uitgeput ben, "
            r"krijg ik die gedachten niet uitgezet. Ik merk ook dat ik minder plezier beleef aan "
            r"sociale momenten. Ik ga nog wel naar afspraken met vrienden, maar ik voel me er "
            r"afstandelijk en niet echt betrokken. Mijn nek en schouders staan bovendien bijna "
            r"voortdurend gespannen en pijnlijk, wat volgens mijn huisarts waarschijnlijk "
            r"stressgerelateerd is.''"
        ),
        criteria=[
            "Inslaapproblemen",
            "Cognitieve hyperarousal",
            "Sociale anhedonie",
            "Stressgerelateerde spierspanning",
        ],
        monitoring=(
            r"Schermtijd na 22.00 uur: 16/21 avonden (gem. 78 min). Openstaande taken: "
            r"gem. 6,2 per dag. Sociale activiteiten: 2,1/week (beleving 2,3/10). "
            r"Preslaap offloading toegepast: 3/21 avonden. Lichamelijke activiteit: "
            r"gem. 0,7 sessies/week."
        ),
        predictors=[
            "Late avondschermtijd",
            "Dagelijkse taakaccumulatie",
            "Preslaap cognitieve offloading",
            "Sociale participatie",
            "Aerobe beweging",
        ],
        tikz_edges=[
            ("cr1", "p1", "S"),
            ("cr2", "p1", "S"),
            ("cr2", "p2", "M"),
            ("cr2", "p3", "S"),
            ("cr1", "p3", "M"),
            ("cr3", "p4", "S"),
            ("cr1", "p5", "M"),
            ("cr4", "p5", "M"),
        ],
        treatment_targets=[
            (
                "Schermtijd na 22.00 uur",
                "hoogste impact op CR-1 en CR-2; monitoring toont consequent laat gebruik "
                "(16/21 avonden)",
            ),
            (
                "Preslaap cognitieve offloading",
                "nauwelijks ingezet ondanks duidelijke taakaccumulatie; rechtstreeks relevant "
                "voor CR-2",
            ),
            (
                "Aerobe beweging",
                "zeer lage frequentie (0,7 sessies/week); potentieel effect op CR-1 en CR-4",
            ),
        ],
        p5_challenge=(
            "Moeizaam inslapen door een overactieve geest en aanhoudende spanning in de avond."
        ),
        p5_target=(
            "Schermtijd in de late avond reduceren en een vaste afschakelroutine installeren."
        ),
        p5_barrier=(
            "Lage zelfeffectiviteit: schermgebruik voelt als de enige haalbare manier om na het "
            "werk te decomprimeren."
        ),
        p5_coping=(
            "Implementatie-intentie met substitutie: vervang de laatste 30 minuten schermtijd "
            "door een korte schrijf- of bewegingsroutine voor het slapengaan."
        ),
        bfs_items=[
            "Cafeine-inname na 16.00 uur (aantal)",
            "Schermtijd na 22:00 uur (min)",
            "Aantal dutjes overdag (min)",
            "Preslaap cognitieve offloading toegepast (ja/nee)",
            "Tijd in bed voor het opstaan (min)",
            "Aerobe beweging vandaag (min)",
            "Aantal sociale afspraken gepland (aantal)",
            "Nek- en schouderspanning in de namiddag (0--10)",
            "Blootstelling aan blauw licht voor het slapengaan (min)",
            "Piekeren over sociale afwijzing (0--10)",
            "Aantal maaltijden na 21.00 uur (aantal)",
            "Ochtendlichtblootstelling (min)",
            "Preslaap ontspanningsritueel uitgevoerd (ja/nee)",
            "Werkuren vandaag (uren)",
            "Alcoholinname in de avond (aantal)",
            "Slaapduur afgelopen nacht (uren)",
            "Aantal onbeantwoorde berichten (aantal)",
            "Pijnintensiteit lage rug (0--10)",
            "Hydratatie vandaag (glazen)",
            "Concentratie overdag (0--10)",
        ],
        bfs_correct=[2, 4, 6, 9, 13],
    ),
    "C02": dict(
        label="C02",
        profile="34-jarige administratief medewerker in de zorg",
        duration="3 maanden",
        short_desc=(
            "Recidiverende paniekepisoden, anticipatieangst, situationele vermijding en "
            "beroepsmatige interferentie."
        ),
        vignette=(
            r"``Ongeveer drie maanden geleden begon ik plots episodes te krijgen waarbij mijn hart "
            r"tekeergaat, ik het gevoel heb dat ik geen lucht krijg en ik er zeker van ben dat er "
            r"iets ernstig misloopt. Die episodes zijn heel beangstigend en kunnen zelfs optreden "
            r"wanneer ik thuis rustig neerzit. Sindsdien ben ik bang geworden om zo'n episode in "
            r"het openbaar te krijgen en vermijd ik supermarkten, openbaar vervoer en drukke "
            r"plaatsen. Ik heb mij al meerdere keren ziek gemeld omdat ik bang was dat het op het "
            r"werk zou gebeuren. Ik weet rationeel dat die situaties niet gevaarlijk zijn, maar de "
            r"angst blijft overheersen.''"
        ),
        criteria=[
            "Recidiverende paniekepisoden",
            "Anticipatieangst",
            "Situationele vermijding",
            "Beroepsmatige interferentie",
        ],
        monitoring=(
            r"Paniekepisoden: gem. 2,3/week. Vermijdingsepisodes: gem. 4,1/dag. "
            r"Veiligheidsgedrag: gem. 8,7/dag. Geplande exposurepogingen: gem. 0,3/dag. "
            r"Werkaanwezigheid: 60\%. Verwachte angst in publieke settings: 7,8/10."
        ),
        predictors=[
            "Vermijdingsgedrag",
            "Veiligheidsgedrag",
            "Interoceptieve focus",
            "Geplande exposure",
            "Dagstructuur",
        ],
        tikz_edges=[
            ("cr3", "p1", "S"),
            ("cr2", "p1", "S"),
            ("cr2", "p2", "S"),
            ("cr1", "p2", "M"),
            ("cr2", "p3", "S"),
            ("cr1", "p3", "S"),
            ("cr3", "p4", "S"),
            ("cr4", "p5", "M"),
        ],
        treatment_targets=[
            (
                "Geplande exposure",
                "nagenoeg afwezig in de monitoring (0,3/dag); rechtstreeks aangrijpingspunt voor "
                "vermijdingsreductie",
            ),
            (
                "Veiligheidsgedrag",
                "hoge frequentie (8,7/dag) onderhoudt anticipatieangst en paniekverwachting",
            ),
            (
                "Interoceptieve focus",
                "verhoogde aandacht voor lichamelijke sensaties voedt escalatie van CR-1 en CR-2",
            ),
        ],
        p5_challenge=(
            "Paniekangst leidt tot toenemende vermijding en duidelijke hinder in het dagelijks functioneren."
        ),
        p5_target=(
            "Geplande exposure aan gevreesde situaties in een hanteerbare, stapsgewijze opbouw."
        ),
        p5_barrier=(
            "Sterke negatieve uitkomstverwachting: de persoon verwacht dat betreden van de situatie "
            "onvermijdelijk tot een onbeheersbare paniekreactie zal leiden."
        ),
        p5_coping=(
            "Werk met een concrete exposurehiërarchie en een vooraf afgesproken copingzin per stap."
        ),
        bfs_items=[
            "Cafeine-inname voor 12.00 uur (aantal)",
            "Geplande exposure-oefening uitgevoerd (ja/nee)",
            "Rusthartslag bij ontwaken (slagen/min)",
            "Spierspanning in schouders (0--10)",
            "Veiligheidsgedrag frequentie (0--10)",
            "Sociale contacten vandaag (aantal)",
            "Lichamelijke sensaties zonder reactie geobserveerd (min)",
            "Stress op het werk (0--10)",
            "Inslaaptijd afgelopen nacht (min)",
            "Vermijdingssituatie betreden ondanks angst (ja/nee)",
            "Waterinname vandaag (glazen)",
            "Plezier in vrijetijdsactiviteit (0--10)",
            "Aantal dutjes overdag (min)",
            "Eetlust vandaag (0--10)",
            "Anticipatieangst voor gevreesde situatie (0--10)",
            "Stemmingsval in de namiddag (0--10)",
            "Stapelteller vandaag (aantal)",
            "Concentratie tijdens administratie (0--10)",
            "Alcoholinname vanavond (aantal)",
            "Tijd buitenshuis zonder ondersteuning (min)",
        ],
        bfs_correct=[2, 5, 7, 10, 15],
    ),
    "C03": dict(
        label="C03",
        profile="41-jarige leraar secundair onderwijs",
        duration="7 maanden",
        short_desc=(
            "Emotionele uitputting, professionele depersonalisatie, sociale terugtrekking en "
            "cognitieve fouten op het werk."
        ),
        vignette=(
            r"``Ik geef al vijftien jaar les en deed dat altijd graag, maar de voorbije maanden voel "
            r"ik mij volledig leeg. Ik ga naar school, geef mijn lessen en kom thuis, maar er is "
            r"niets meer over. 's Avonds zit ik voor de televisie en kan ik mij nergens toe "
            r"aanzetten. Ik ben ook gestopt met afspreken met vrienden en reageer nog zelden op "
            r"berichten. Op het werk begin ik fouten te maken -- dingen vergeten, de draad kwijt "
            r"raken tijdens een les -- wat vroeger nooit gebeurde. Ik schaam mij daarvoor, maar ik "
            r"weet niet hoe ik dit moet keren.''"
        ),
        criteria=[
            "Emotionele uitputting",
            "Professionele depersonalisatie",
            "Sociale terugtrekking",
            "Cognitieve fouten op het werk",
        ],
        monitoring=(
            r"Werkoverschrijding na werktijd: gem. 5,4 episodes/dag. Herstelactiviteiten: "
            r"gem. 1,2/week. Sociaal contact op eigen initiatief: gem. 0,7/week. "
            r"Professionele meestervaringen: gem. 0,9/dag. Slaapkwaliteit: gem. 4,1/10."
        ),
        predictors=[
            "Werk-privegrens",
            "Herstelactiviteit",
            "Sociaal initiatief",
            "Professionele meestervaring",
            "Slaapkwaliteit",
        ],
        tikz_edges=[
            ("cr1", "p1", "S"),
            ("cr2", "p1", "S"),
            ("cr1", "p2", "S"),
            ("cr3", "p2", "M"),
            ("cr3", "p3", "S"),
            ("cr2", "p4", "M"),
            ("cr4", "p4", "M"),
            ("cr1", "p5", "M"),
            ("cr4", "p5", "M"),
        ],
        treatment_targets=[
            (
                "Werk-privegrens",
                "sterkste onderhoudende factor voor uitputting; frequente overschrijding na werktijd",
            ),
            (
                "Herstelactiviteit",
                "kritisch laag niveau (1,2/week); relevant voor CR-1 en CR-3",
            ),
            (
                "Professionele meestervaring",
                "kan depersonalisatie temperen en dag-tot-dag positieve variatie benutten",
            ),
        ],
        p5_challenge=(
            "Uitgesproken werkgerelateerde uitputting met verlies aan betekenis en toenemende sociale terugtrekking."
        ),
        p5_target=(
            "Een duidelijke werk-privegrens installeren en na het werk consequent herstelgedrag activeren."
        ),
        p5_barrier=(
            "Actieplanning ontbreekt: er is geen concreet moment of signaal waarop de werkmodus stopt."
        ),
        p5_coping=(
            "Koppel een vaste afsluitsignalering aan een kleine herstelactiviteit die onmiddellijk na "
            "het einde van de werkdag start."
        ),
        bfs_items=[
            "Aantal huishoudelijke taken voltooid (aantal)",
            "Werkoverschrijding na werktijd (min)",
            "Piekeren over gezinszaken in de avond (0--10)",
            "Cafeine-inname na 15.00 uur (aantal)",
            "Herstelactiviteit vandaag uitgevoerd (ja/nee)",
            "Aantal lesuren vandaag (uren)",
            "Eetlust bij avondmaal (0--10)",
            "Professionele meestervaring vandaag (0--10)",
            "Ochtendspanning bij ontwaken (0--10)",
            "Aantal nieuwschecks in de avond (aantal)",
            "Slaapkwaliteit bij ontwaken (0--10)",
            "Lichamelijke pijnklachten na het werk (0--10)",
            "Administratieve achterstand (0--10)",
            "Reistijd woon-werk (min)",
            "Middagdutje genomen (ja/nee)",
            "Sociaal contact op eigen initiatief (ja/nee)",
            "Aantal maaltijden overgeslagen (aantal)",
            "Schermtijd in de late avond (min)",
            "Irritatie tegenover collega's (0--10)",
            "Weekendplanning opgesteld (ja/nee)",
        ],
        bfs_correct=[2, 5, 8, 11, 16],
    ),
    "C04": dict(
        label="C04",
        profile="63-jarige gepensioneerde professional",
        duration="9 maanden na overlijden van partner",
        short_desc=(
            "Aanhoudende rouwreactie, doorslaapinsomnie, schuldgedreven ruminatie en verminderde "
            "toekomstorientatie."
        ),
        vignette=(
            r"``Mijn partner is negen maanden geleden overleden en ik geraak daar nog steeds niet door. "
            r"Ik denk voortdurend aan hem of haar. Kleine prikkels -- een liedje, een geur, een "
            r"televisieprogramma dat we samen bekeken -- lokken intense golven van verdriet uit die "
            r"ik niet onder controle krijg. Ik slaap slecht; ik word vroeg wakker en raak daarna "
            r"niet meer in slaap. Daarnaast voel ik veel schuld. Ik blijf maar herhalen wat ik had "
            r"moeten zeggen of anders had moeten doen. Ik kan mij ook nauwelijks voorstellen hoe mijn "
            r"toekomst er nog betekenisvol kan uitzien. Mijn familie zegt dat het al beter zou moeten "
            r"gaan, maar zo voelt het niet.''"
        ),
        criteria=[
            "Aanhoudende rouwreactie",
            "Doorslaapinsomnie",
            "Schuldgedreven ruminatie",
            "Verminderde toekomstorientatie",
        ],
        monitoring=(
            r"Rouwintrusies: gem. 4,2/dag. Slaapkwaliteit: gem. 2,8/10. Sociale contactmomenten: "
            r"gem. 3,1/week. Schuldgedachten: gem. 5,8/dag. Toekomstgerichte activiteiten: "
            r"gem. 0,4/week. Betekenisgerichte activiteit uitgevoerd: 4/21 dagen."
        ),
        predictors=[
            "Betekenisgerichte activiteit",
            "Slaap-waakconsistentie",
            "Sociale steunopname",
            "Zelfcompassie",
            "Toekomstplanning",
        ],
        tikz_edges=[
            ("cr1", "p1", "S"),
            ("cr4", "p1", "S"),
            ("cr2", "p2", "S"),
            ("cr1", "p3", "M"),
            ("cr3", "p3", "M"),
            ("cr3", "p4", "S"),
            ("cr4", "p5", "S"),
        ],
        treatment_targets=[
            (
                "Betekenisgerichte activiteit",
                "slechts 4/21 dagen uitgevoerd; hoog potentieel voor CR-1 en CR-4",
            ),
            (
                "Slaap-waakconsistentie",
                "lage slaapkwaliteit ondermijnt emotieregulatie en rouwverwerking",
            ),
            (
                "Zelfcompassie",
                "schuldgedreven ruminatie is hoog en vraagt een direct corrigerend proces",
            ),
        ],
        p5_challenge=(
            "Aanhoudende rouw met uitgesproken schuldbeleving en weinig toekomstgerichte betrokkenheid."
        ),
        p5_target=(
            "Een kleine maar betekenisvolle activiteit plannen die verbinding maakt met herinnering en toekomst."
        ),
        p5_barrier=(
            "Motivationele inertie: de eerste stap voelt zinloos of emotioneel te zwaar."
        ),
        p5_coping=(
            "Verklein de instap naar een tijdsafgebakende micro-actie, bijvoorbeeld een korte schrijfoefening "
            "of een vast herinneringsmoment."
        ),
        bfs_items=[
            "Wandeling buiten uitgevoerd (ja/nee)",
            "App-berichten met familie verstuurd (aantal)",
            "Betekenisgerichte activiteit uitgevoerd (ja/nee)",
            "Cafeine-inname na 17.00 uur (aantal)",
            "Televisietijd in de avond (min)",
            "Vaste slaap-waakroutine gevolgd (ja/nee)",
            "Maaltijdlust vandaag (0--10)",
            "Werk- of zorgtaken voor anderen (min)",
            "Zelfcompassie-oefening vandaag uitgevoerd (ja/nee)",
            "Middagdutje genomen (ja/nee)",
            "Lichamelijke onrust in de avond (0--10)",
            "Intrusieve rouwgedachten intensiteit (0--10)",
            "Schermtijd na 21.00 uur (min)",
            "Vermijden van herinneringsvoorwerpen (ja/nee)",
            "Alcoholinname in de avond (aantal)",
            "Huishoudelijke taken voltooid (aantal)",
            "Sociaal contactmoment vandaag (ja/nee)",
            "Aantal afspraken buitenshuis gepland (aantal)",
            "Pijnintensiteit in gewrichten (0--10)",
            "Ochtendenergie bij opstaan (0--10)",
        ],
        bfs_correct=[3, 6, 9, 12, 17],
    ),
    "C05": dict(
        label="C05",
        profile="25-jarige postgraduaatsstudent",
        duration="ongeveer 20 maanden",
        short_desc=(
            "Intrusieve schadeobsessies, compulsieve mentale controle, vermijding van schaderisico "
            "en uitgesproken tijdsbelasting door OCD-klachten."
        ),
        vignette=(
            r"``Ik krijg gedachten dat ik per ongeluk iemand zou kunnen verwonden, ook al wil ik dat "
            r"helemaal niet en weet ik dat die gedachten irrationeel zijn. Ze komen plots op en "
            r"voelen dan heel echt en beangstigend aan. Daarna ben ik uren bezig met in mijn hoofd "
            r"na te gaan of ik niets fout heb gedaan. Ik ben ook situaties beginnen vermijden waarin "
            r"ik iets scherp zou kunnen vasthouden, wat mijn dagelijks leven en sociale contacten "
            r"beinvloedt. Dat mentale controleren neemt uren in beslag en begint mijn studies en "
            r"relaties serieus te verstoren.''"
        ),
        criteria=[
            "Intrusieve schadeobsessies",
            "Compulsieve mentale controle",
            "Vermijding van schaderisico",
            "Tijdsbelasting door OCD",
        ],
        monitoring=(
            r"Compulsie- en controletijd: gem. 2,4 uur/dag. Vermijdingsepisodes: gem. 3,8/dag. "
            r"Weerstandsoefening: gem. 18 min/dag. Distress bij niet reageren: 7,2/10. "
            r"Slaap: gem. 5,6 uur/nacht. Geruststellingsvragen: gem. 4,1/dag."
        ),
        predictors=[
            "Compulsie-uitstel",
            "Vermijdingsbeheer",
            "Onzekerheidstolerantie",
            "Slaapkwaliteit",
            "Geruststellingszoekend gedrag",
        ],
        tikz_edges=[
            ("cr2", "p1", "S"),
            ("cr4", "p1", "S"),
            ("cr3", "p2", "S"),
            ("cr1", "p3", "S"),
            ("cr2", "p3", "S"),
            ("cr1", "p4", "M"),
            ("cr2", "p5", "S"),
            ("cr4", "p5", "M"),
        ],
        treatment_targets=[
            (
                "Compulsie-uitstel",
                "weinig toegepast ten opzichte van de benodigde ERP-intensiteit; centraal aangrijpingspunt",
            ),
            (
                "Vermijdingsbeheer",
                "hoge vermijdingsfrequentie onderhoudt de obsessie-compulsiecyclus",
            ),
            (
                "Onzekerheidstolerantie",
                "lage tolerantie voedt zowel intrusies als mentale controle",
            ),
        ],
        p5_challenge=(
            "Belastende intrusieve gedachten leiden tot langdurige mentale controle en duidelijke functionele hinder."
        ),
        p5_target=(
            "De tijd tussen intrusie en controlehandeling systematisch verlengen."
        ),
        p5_barrier=(
            "Overschat risico: niet controleren voelt moreel onverantwoord en potentieel gevaarlijk."
        ),
        p5_coping=(
            "Gebruik een uitstelregel met een korte grondingsoefening als vast alternatief tijdens de eerste minuten."
        ),
        bfs_items=[
            "Schermtijd na 22.00 uur (min)",
            "Compulsie-uitstelpoging duur (min)",
            "Slaapduur afgelopen nacht (uren)",
            "Spierspanning in de schouders (0--10)",
            "Vermijdingssituatie betreden ondanks obsessie (ja/nee)",
            "Aantal sociale afspraken vandaag (aantal)",
            "Cafeine-inname na de middag (aantal)",
            "Onzekerheidstolerantieoefening uitgevoerd (ja/nee)",
            "Studie-uren vandaag (uren)",
            "Somatische angstsensaties (0--10)",
            "Weerstand geboden aan herhalingsgedachte (0--10)",
            "Lichamelijke activiteit vandaag (min)",
            "Maaltijdlust vandaag (0--10)",
            "Slaapritme consistent gevolgd (ja/nee)",
            "Emotieregulatievaardigheid gebruikt (ja/nee)",
            "Geruststelling gevraagd (aantal)",
            "Alcoholinname in de avond (aantal)",
            "Concentratie tijdens studeren (0--10)",
            "Contact met vrienden vandaag (ja/nee)",
            "Ochtendstemming bij ontwaken (0--10)",
        ],
        bfs_correct=[2, 5, 8, 11, 16],
    ),
    "C06": dict(
        label="C06",
        profile="32-jarige marketingprofessional",
        duration="persisterend sinds het begin van de loopbaan (meer dan 3 jaar)",
        short_desc=(
            "Prestatieangst op het werk, post-event ruminatie, excessieve voorbereiding en "
            "vermijding van zichtbaarheid."
        ),
        vignette=(
            r"``Ik ervaar heel veel angst in professionele situaties. Wanneer ik moet presenteren, "
            r"in een vergadering iets moet zeggen of feedback krijg van mijn leidinggevende, word ik "
            r"zo zelfbewust dat mijn hoofd leegloopt. Vooraf spendeer ik uren aan voorbereiden, maar "
            r"in het moment zelf blokkeer ik toch. Achteraf blijf ik lang analyseren wat ik precies "
            r"heb gezegd en hoe ik overkwam. Ik heb al meerdere promoties en interessante projecten "
            r"afgeslagen omdat ik bang was voor de extra zichtbaarheid. Ik besef dat mijn angst voor "
            r"de beoordeling door anderen buiten proportie is, maar ik krijg ze niet stil.''"
        ),
        criteria=[
            "Prestatieangst op het werk",
            "Post-event ruminatie",
            "Excessieve voorbereiding",
            "Vermijding van zichtbaarheid",
        ],
        monitoring=(
            r"Professionele exposures aangegaan: gem. 0,9/week. Post-event verwerkingstijd: "
            r"gem. 68 min per gebeurtenis. Pre-event voorbereiding: gem. 3,2 uur per gebeurtenis. "
            r"Zelfbeoordeling versus externe feedback: 4,1 versus 6,8/10. "
            r"Afgeslagen doorgroeikansen: 3 in de voorbije 4 weken."
        ),
        predictors=[
            "Professionele exposure",
            "Post-event verwerking",
            "Zelfevaluatiebias",
            "Voorbereidingstijd",
            "Vermijdingslogging",
        ],
        tikz_edges=[
            ("cr1", "p1", "S"),
            ("cr4", "p1", "S"),
            ("cr2", "p2", "S"),
            ("cr1", "p3", "M"),
            ("cr2", "p3", "M"),
            ("cr3", "p4", "S"),
            ("cr4", "p5", "S"),
        ],
        treatment_targets=[
            (
                "Professionele exposure",
                "lage frequentie terwijl dit het kernmechanisme is voor angstreductie",
            ),
            (
                "Post-event verwerking",
                "gemiddeld 68 minuten per gebeurtenis; duidelijke onderhoudende factor voor CR-2",
            ),
            (
                "Vermijdingslogging",
                "zichtbaar maken van gemiste kansen is cruciaal voor CR-4 en behandelsturing",
            ),
        ],
        p5_challenge=(
            "Sterke angst in professionele evaluatie- en zichtbaarheidssituaties met langdurige naverwerking."
        ),
        p5_target=(
            "Professionele exposure uitvoeren ondanks spanning en achteraf functioneel evalueren."
        ),
        p5_barrier=(
            "Grote kloof in zelfeffectiviteit: de persoon verwacht dat exposure zal bevestigen dat hij of zij tekortschiet."
        ),
        p5_coping=(
            "Gebruik een gedragsproef waarbij voorspelde prestatie expliciet wordt vergeleken met feitelijke feedback."
        ),
        bfs_items=[
            "Slaapduur afgelopen nacht (uren)",
            "Aantal koffies voor een meeting (aantal)",
            "Professionele exposure-oefening uitgevoerd (ja/nee)",
            "Lichamelijke spanning voor het werk (0--10)",
            "Post-event verwerkingstijd (min)",
            "Werkuren vandaag (uren)",
            "Positieve feedback ontvangen (ja/nee)",
            "Vermeden kans gedocumenteerd (ja/nee)",
            "Aantal sociale berichten verstuurd (aantal)",
            "Middagpauze genomen (ja/nee)",
            "Schermtijd na 21.00 uur (min)",
            "Zelfevaluatie vergeleken met externe feedback (ja/nee)",
            "Eetlust voor de lunch (0--10)",
            "Aerobe beweging vandaag (min)",
            "Irritatie naar collega's (0--10)",
            "Pre-event voorbereidingstijd (min)",
            "Alcoholinname na het werk (aantal)",
            "Ochtendmoeheid (0--10)",
            "Aantal huishoudelijke taken afgerond (aantal)",
            "Sociale steun gevraagd (ja/nee)",
        ],
        bfs_correct=[3, 5, 8, 12, 16],
    ),
    "C07": dict(
        label="C07",
        profile="46-jarige accountmanager",
        duration="14 maanden",
        short_desc=(
            "Aanhoudende sombere stemming, anergie, sociale isolatie en hopeloosheid over herstel."
        ),
        vignette=(
            r"``Al meer dan een jaar word ik de meeste ochtenden wakker met een zwaar en vlak gevoel. "
            r"Het is niet eens intense verdrietigheid, eerder een aanhoudende grijsheid waardoor "
            r"niets nog de moeite lijkt. Heel eenvoudige taken voelen enorm zwaar aan: mij aankleden, "
            r"een mail beantwoorden, een maaltijd maken. Ik ben stilaan het contact verloren met bijna "
            r"iedereen met wie ik vroeger omging, omdat ik de energie niet heb om relaties te onderhouden. "
            r"Ik slaap veel maar voel mij nooit uitgerust. In stilte vraag ik mij af of ik ooit nog de "
            r"oude zal worden. Mijn huisarts heeft mij aangeraden hulp te zoeken.''"
        ),
        criteria=[
            "Aanhoudende sombere stemming",
            "Anergie",
            "Sociale isolatie",
            "Hopeloosheid",
        ],
        monitoring=(
            r"Voltooiingsgraad van geplande activiteiten: 28\%. Sociale contactmomenten/week: "
            r"gem. 0,8. Stemming: gem. 2,9/10. Variatie in slaap-waaktijd: gem. 2,3 uur. "
            r"Plezierige activiteiten: gem. 0,9/dag. Uitdagen van hopeloze gedachten: gem. 0,6/dag."
        ),
        predictors=[
            "Activiteitenschema",
            "Sociaal initiatief",
            "Slaapregelmaat",
            "Plezierige activiteit",
            "Cognitieve tegenspraak hopeloosheid",
        ],
        tikz_edges=[
            ("cr1", "p1", "S"),
            ("cr2", "p1", "S"),
            ("cr3", "p2", "S"),
            ("cr1", "p3", "M"),
            ("cr2", "p3", "M"),
            ("cr1", "p4", "S"),
            ("cr4", "p4", "M"),
            ("cr4", "p5", "S"),
        ],
        treatment_targets=[
            (
                "Activiteitenschema",
                r"lage voltooiingsgraad (28\%); gedragsactivatie is hier het primaire mechanisme",
            ),
            (
                "Slaapregelmaat",
                "sterke variatie in slaap-waakritme verstoort stemming en energieregulatie",
            ),
            (
                "Sociaal initiatief",
                "gemiddeld 0,8 contacten/week; essentieel voor CR-3 en indirect voor stemming",
            ),
        ],
        p5_challenge=(
            "Langdurige sombere stemming en anergie hebben activiteit en sociale verbinding sterk doen afnemen."
        ),
        p5_target=(
            "Dagelijks een klein, gepland gedrag uitvoeren als onderdeel van gedragsactivatie."
        ),
        p5_barrier=(
            "Planningsdrempel en betekenisverlies: activiteiten voelen vooraf nutteloos en te zwaar."
        ),
        p5_coping=(
            "Werk met een vooraf gekozen micro-activiteit en een 'klaar is genoeg'-criterium."
        ),
        bfs_items=[
            "Cafeine-inname na 15.00 uur (aantal)",
            "Activiteitenschema gevolgd (ja/nee)",
            "Eetlust in de ochtend (0--10)",
            "Werkgerelateerd piekeren in de avond (min)",
            "Slaap-waaktijdvariatie vandaag (min)",
            "Pijnklachten in de rug (0--10)",
            "Sociaal contact geinitieerd (ja/nee)",
            "Tijd op sociale media (min)",
            "Middagdutje genomen (ja/nee)",
            "Plezierige activiteit voltooid (ja/nee)",
            "Alcoholinname vanavond (aantal)",
            "Concentratie tijdens administratie (0--10)",
            "Aantal maaltijden overgeslagen (aantal)",
            "Hopeloosheidsgedachten uitgedaagd (ja/nee)",
            "Waterinname vandaag (glazen)",
            "Werkuren vandaag (uren)",
            "TV-kijktijd in de avond (min)",
            "Schermtijd na 22.00 uur (min)",
            "Aantal afspraken buitenshuis gepland (aantal)",
            "Ochtendspanning (0--10)",
        ],
        bfs_correct=[2, 5, 7, 10, 14],
    ),
    "C08": dict(
        label="C08",
        profile="38-jarige projectmanager",
        duration="sinds de kindertijd, duidelijk verergerd in de voorbije 12 maanden",
        short_desc=(
            "Aandachtsvolhoudingsproblemen, plannings- en organisatiefouten, interpersoonlijke "
            "onoplettendheid en prestatieschaamte."
        ),
        vignette=(
            r"``Ik heb altijd al moeite gehad om mij te concentreren op taken die mij niet meteen "
            r"interesseren, maar het is het voorbije jaar veel erger geworden. Ik begin aan projecten "
            r"en maak ze zelden af. Ik vergeet afspraken, zelfs als ik ze ergens noteer. Mijn partner "
            r"is gefrustreerd omdat ik volgens hem of haar niet echt luister en onbetrouwbaar overkom. "
            r"Ik weet dat ik op veel vlakken competent ben, maar ik voel mij voortdurend een mislukking. "
            r"Ik ben taken die moeilijk zullen zijn ook steeds meer beginnen vermijden, waardoor de "
            r"achterstand alleen maar groter wordt.''"
        ),
        criteria=[
            "Aandachtsvolhoudingsproblemen",
            "Plannings- en organisatiefouten",
            "Interpersoonlijke onoplettendheid",
            "Prestatieschaamte",
        ],
        monitoring=(
            r"Gestarte versus afgewerkte taken: 68\% versus 29\%. Gemiste afspraken of toezeggingen: "
            r"gem. 2,1/week. Partnerconflicten: gem. 4,7/week. Externe tools gebruikt: gem. 1,8/dag. "
            r"Positieve zelfbekrachtiging: gem. 0,4/dag."
        ),
        predictors=[
            "Extern taakbeheersysteem",
            "Interesse-gedreven activatie",
            "Bewust luisteren",
            "Taakvoltooiingsbekrachtiging",
            "Overgangsrituelen",
        ],
        tikz_edges=[
            ("cr2", "p1", "S"),
            ("cr1", "p1", "S"),
            ("cr1", "p2", "M"),
            ("cr3", "p3", "S"),
            ("cr4", "p4", "S"),
            ("cr2", "p4", "M"),
            ("cr2", "p5", "S"),
        ],
        treatment_targets=[
            (
                "Extern taakbeheersysteem",
                "inconsistent gebruik (1,8/dag) terwijl dit de kernscaffolding vormt voor CR-1 en CR-2",
            ),
            (
                "Bewust luisteren",
                "partnerconflicten zijn frequent; direct relevant voor CR-3",
            ),
            (
                "Taakvoltooiingsbekrachtiging",
                "bijna afwezig; kan de schaamtecyclus en vermijding doorbreken",
            ),
        ],
        p5_challenge=(
            "Executieve functieproblemen zorgen voor taakfalen, relationele spanning en aanhoudende schaamte."
        ),
        p5_target=(
            "Een extern taakbeheersysteem consequent raadplegen en successen expliciet bekrachtigen."
        ),
        p5_barrier=(
            "Gewoonteweerstand en lage volhoudingsverwachting: eerdere systemen werden snel losgelaten."
        ),
        p5_coping=(
            "Installeer een minimale, dagelijks terugkerende check-in en koppel afgeronde taken aan een expliciete bevestiging."
        ),
        bfs_items=[
            "Cafeine-inname voor de middag (aantal)",
            "Extern taakbeheersysteem geraadpleegd (ja/nee)",
            "Slaapduur afgelopen nacht (uren)",
            "Tijd op sociale media tijdens werk (min)",
            "Spierspanning in de namiddag (0--10)",
            "Bewust luisteren geoefend (ja/nee)",
            "Ochtendlichtblootstelling (min)",
            "Aantal onderbrekingen door notificaties (aantal)",
            "Taakvoltooiing erkend of beloond (ja/nee)",
            "Emotionele kwetsbaarheid na kritiek (0--10)",
            "Aantal huishoudelijke taken gestart (aantal)",
            "Eetlust rond de lunch (0--10)",
            "Overgangstaak-ritueel gebruikt (ja/nee)",
            "Alcoholinname in de avond (aantal)",
            "Stemmingsschommeling vandaag (0--10)",
            "Stapelteller vandaag (aantal)",
            "Gemaakte toezeggingen nagekomen (0--10)",
            "Middagdutje genomen (ja/nee)",
            "Schermtijd na 22.00 uur (min)",
            "Plezier in vrijetijdsbesteding (0--10)",
        ],
        bfs_correct=[2, 6, 9, 13, 17],
    ),
    "C09": dict(
        label="C09",
        profile="27-jarige salesvertegenwoordiger",
        duration="3 jaar",
        short_desc=(
            "Snelle stemmingsschommelingen, uitgesproken afwijzingsreactiviteit, impulsief gedrag "
            "en relationele instabiliteit."
        ),
        vignette=(
            r"``Mijn stemming kan op een dag heel snel omslaan en vaak lijkt dat sterker dan wat er "
            r"feitelijk gebeurt. Ik kan mij in de ochtend nog redelijk voelen en tegen de middag "
            r"helemaal ontredderd zijn. Wanneer ik mij bekritiseerd of afgewezen voel, zelfs om iets "
            r"kleins, reageer ik extreem heftig -- boos of heel wanhopig. Ik handel ook impulsief, "
            r"vooral met geld of in relaties, en heb daar achteraf spijt van. Mijn relaties worden "
            r"vaak snel intens en lopen daarna mis. Ik zie dat dat patroon zich herhaalt, maar ik weet "
            r"niet hoe ik het moet doorbreken.''"
        ),
        criteria=[
            "Snelle stemmingsschommelingen",
            "Afwijzingsreactiviteit",
            "Impulsief gedrag",
            "Relationele instabiliteit",
        ],
        monitoring=(
            r"Dagelijks stembereik: gem. 5,8 punten. Emotieregulatievaardigheden gebruikt: "
            r"gem. 1,2/dag. Impulsieve handelingen: gem. 2,1/dag. Interpersoonlijke conflicten: "
            r"gem. 3,8/week. Ervaren afwijzingssignalen: gem. 2,4/dag. Herwaardering voor reactie: "
            r"gem. 0,7/dag."
        ),
        predictors=[
            "Emotieregulatievaardigheid",
            "Afwijzingsherwaardering",
            "Impulsuitstel",
            "Interpersoonlijke communicatie",
            "Gedifferentieerde stemmingsmonitoring",
        ],
        tikz_edges=[
            ("cr1", "p1", "S"),
            ("cr2", "p1", "S"),
            ("cr2", "p2", "S"),
            ("cr4", "p2", "S"),
            ("cr3", "p3", "S"),
            ("cr4", "p4", "S"),
            ("cr2", "p4", "M"),
            ("cr1", "p5", "M"),
        ],
        treatment_targets=[
            (
                "Emotieregulatievaardigheid",
                "laag gebruik ondanks groot stembereik; meest directe hefboom voor stabilisatie",
            ),
            (
                "Impulsuitstel",
                "impulsieve handelingen komen frequent voor en voeden relationele schade",
            ),
            (
                "Afwijzingsherwaardering",
                "grootste kloof tussen ervaren afwijzing en voorafgaande cognitieve bijsturing",
            ),
        ],
        p5_challenge=(
            "Snelle stemmingswisselingen en hevige reactiviteit leiden tot impulsieve beslissingen en relationele escalatie."
        ),
        p5_target=(
            "Een specifieke emotieregulatievaardigheid inzetten zodra spanning begint op te lopen."
        ),
        p5_barrier=(
            "Lage coping-self-efficacy: in het piekmoment verwacht de persoon dat vaardigheden toch niet zullen werken."
        ),
        p5_coping=(
            "Gebruik een vooraf geoefende cue-actie-koppeling: herken een vroege lichamelijke cue en koppel die aan exact een vaardigheid."
        ),
        bfs_items=[
            "Cafeine-inname na 14.00 uur (aantal)",
            "Slaapduur afgelopen nacht (uren)",
            "Emotieregulatievaardigheid toegepast (ja/nee)",
            "Aantal conflicten op het werk (aantal)",
            "Tijd op sociale media (min)",
            "Impulshandeling uitgesteld (ja/nee)",
            "Eetlust bij avondmaal (0--10)",
            "Ochtendenergie bij ontwaken (0--10)",
            "Afwijzingsherwaardering geprobeerd (ja/nee)",
            "Alcoholinname vandaag (aantal)",
            "Lichamelijke activiteit vandaag (min)",
            "Stemmingsbereik vandaag (0--10)",
            "Aantal onafgewerkte taken (aantal)",
            "Concentratie tijdens administratie (0--10)",
            "Middagdutje genomen (ja/nee)",
            "Schermtijd na 22.00 uur (min)",
            "Interpersoonlijk communicatiegesprek geoefend (ja/nee)",
            "Cafeine-inname voor een afspraak (aantal)",
            "Tijd buitenshuis alleen (min)",
            "Nekspanning in de namiddag (0--10)",
        ],
        bfs_correct=[3, 6, 9, 12, 17],
    ),
    "C10": dict(
        label="C10",
        profile="52-jarige operationeel directeur",
        duration="ongeveer 9 maanden",
        short_desc=(
            "Chronische werkstress, doorslaapinsomnie, somatische stressklachten en alcoholgebruik "
            "als coping."
        ),
        vignette=(
            r"``De werkstress is het grootste deel van het afgelopen jaar blijven oplopen. Ik heb "
            r"voortdurend spanningshoofdpijn en mijn maag is op de meeste dagen ontregeld. 's Nachts "
            r"lig ik wakker over werkproblemen en beslissingen die nog genomen moeten worden, en als "
            r"ik eenmaal wakker ben geraakt, val ik vaak niet meer in slaap. Ik drink bijna elke avond "
            r"twee of drie glazen wijn om tot rust te komen. Ik weet dat dat niet ideaal is, maar het "
            r"voelt alsof ik anders niet kan afschakelen. Mijn beweging is volledig weggevallen en ik "
            r"snap thuis sneller dan vroeger.''"
        ),
        criteria=[
            "Chronische werkstress",
            "Doorslaapinsomnie",
            "Somatische stressklachten",
            "Alcohol als copingstrategie",
        ],
        monitoring=(
            r"Werkgerelateerde gedachten na 20.00 uur: gem. 3,4 uur. Alcoholeenheden in de avond: "
            r"gem. 3,1. Hoofdpijndagen: 16/21. Slaapefficientie: gem. 62\%. "
            r"Stressregulatie overdag: gem. 0,9/dag. Lichamelijke activiteit: gem. 0,4 sessies/week."
        ),
        predictors=[
            "Werkontkoppeling voor slaap",
            "Avondlijk alcoholgebruik",
            "Stressregulatie overdag",
            "Aerobe beweging",
            "Nachtelijke cognitieve arousal",
        ],
        tikz_edges=[
            ("cr2", "p1", "S"),
            ("cr1", "p1", "S"),
            ("cr4", "p2", "S"),
            ("cr2", "p2", "S"),
            ("cr1", "p3", "M"),
            ("cr3", "p3", "S"),
            ("cr1", "p4", "M"),
            ("cr3", "p4", "M"),
            ("cr2", "p5", "S"),
        ],
        treatment_targets=[
            (
                "Werkontkoppeling voor slaap",
                "momenteel afwezig en direct relevant voor doorslaapinsomnie",
            ),
            (
                "Avondlijk alcoholgebruik",
                "gemiddeld 3,1 eenheden per avond; onderhoudt slaapfragmentatie",
            ),
            (
                "Stressregulatie overdag",
                "laag maar trainbaar; upstream aangrijpingspunt voor CR-1 en CR-3",
            ),
        ],
        p5_challenge=(
            "Aanhoudende werkstress leidt tot gefragmenteerde slaap, lichamelijke spanningsklachten en alcohol als vaste ontladingsstrategie."
        ),
        p5_target=(
            "Een korte werkontkoppelingsroutine installeren en alcohol niet langer als standaard afschakelmechanisme gebruiken."
        ),
        p5_barrier=(
            "Sterke gewoontekracht: alcohol voelt essentieel om de overgang van werk naar slaap te maken."
        ),
        p5_coping=(
            "Vervang het bestaande gewoontepatroon door een vast, tijdsgebonden afbouwritueel dat elke avond op hetzelfde moment start."
        ),
        bfs_items=[
            "Ochtendlichtblootstelling (min)",
            "Preslaap werkontkoppelingsritueel gevolgd (ja/nee)",
            "Middagdutje genomen (ja/nee)",
            "Aantal sociale afspraken na het werk (aantal)",
            "Alcoholeenheden vanavond (aantal)",
            "Eetlust bij avondmaal (0--10)",
            "Hoofdpijnintensiteit in de namiddag (0--10)",
            "Stressregulatiegedrag overdag toegepast (ja/nee)",
            "Schermtijd voor ontspanning na 21.00 uur (min)",
            "Lichamelijke spierspanning in de ochtend (0--10)",
            "Waterinname vandaag (glazen)",
            "Werkgerelateerde gedachten na 20.00 uur (min)",
            "Aantal vergaderuren vandaag (uren)",
            "Slaapduur afgelopen nacht (uren)",
            "TV-kijktijd in de avond (min)",
            "Ochtendstemming bij ontwaken (0--10)",
            "Aerobe beweging vandaag (min)",
            "Cafeine-inname na 16.00 uur (aantal)",
            "Contact met gezin vandaag (ja/nee)",
            "Aantal maaltijden overgeslagen (aantal)",
        ],
        bfs_correct=[2, 5, 8, 12, 17],
    ),
}


ASSIGNMENT = {
    1: ("HCP-PRE-01", "C01", "C02"),
    2: ("HCP-PRE-02", "C03", "C04"),
    3: ("HCP-PRE-03", "C05", "C06"),
    4: ("HCP-PRE-04", "C07", "C08"),
    5: ("HCP-PRE-05", "C09", "C10"),
}


def J(*parts: str) -> str:
    return "\n".join(parts)


def validate_cases() -> None:
    for case_id, case in CASES.items():
        assert len(case["criteria"]) == 4, f"{case_id}: verwacht 4 criteria"
        assert len(case["predictors"]) == 5, f"{case_id}: verwacht 5 predictors"
        assert len(case["treatment_targets"]) == 3, f"{case_id}: verwacht 3 behandeldoelen"
        assert len(case["bfs_items"]) == 20, f"{case_id}: verwacht 20 BFS-items"
        assert len(case["bfs_correct"]) == 5, f"{case_id}: verwacht 5 correcte BFS-posities"
        assert len(set(case["bfs_correct"])) == 5, f"{case_id}: dubbele BFS-sleutels"
        assert all(1 <= idx <= 20 for idx in case["bfs_correct"]), f"{case_id}: BFS-index buiten bereik"


validate_cases()


PREAMBLE_TEMPLATE = r"""\documentclass[11pt,a4paper]{article}

\usepackage[a4paper,left=2.2cm,right=2.2cm,top=2.4cm,bottom=2.4cm]{geometry}
\usepackage[T1]{fontenc}
\usepackage[utf8]{inputenc}
\usepackage{lmodern}
\usepackage{microtype}
\usepackage{booktabs}
\usepackage{tabularx}
\usepackage{array}
\usepackage{enumitem}
\usepackage{hyperref}
\usepackage{xcolor}
\usepackage{titlesec}
\usepackage{fancyhdr}
\usepackage{lastpage}
\usepackage{tcolorbox}
\usepackage{amsmath}
\usepackage{amssymb}
\usepackage{tikz}
\usetikzlibrary{positioning}
\tcbuselibrary{skins,breakable}

\hypersetup{
  colorlinks=true,
  linkcolor=black,
  urlcolor=black,
  pdftitle={<<PDFTITLE>>},
  pdfauthor={Stijn Van Severen -- Universiteit Gent}
}

\definecolor{PrimaryBlue}{HTML}{1D4ED8}
\definecolor{DarkBlue}{HTML}{1E3A5F}
\definecolor{SlateMid}{HTML}{475569}
\definecolor{ForestGreen}{HTML}{047857}
\definecolor{RichPurple}{HTML}{6D28D9}
\definecolor{GoldAmber}{HTML}{B45309}
\definecolor{StrongRed}{HTML}{B91C1C}
\definecolor{SoftBG}{HTML}{F8FAFC}
\definecolor{BorderGrey}{HTML}{CBD5E1}
\definecolor{AccentBlue}{HTML}{DBEAFE}
\definecolor{AccentGreen}{HTML}{D1FAE5}
\definecolor{AccentAmber}{HTML}{FEF3C7}
\definecolor{AccentPurple}{HTML}{EDE9FE}
\definecolor{AccentTeal}{HTML}{CCFBF1}

\titleformat{\section}{\Large\bfseries\color{DarkBlue}}{\thesection}{0.7em}{}[\vspace{-0.2em}\color{PrimaryBlue}\rule{\textwidth}{0.6pt}]
\titleformat{\subsection}{\large\bfseries\color{PrimaryBlue}}{\thesubsection}{0.6em}{}
\titleformat{\subsubsection}{\normalsize\bfseries\color{ForestGreen}}{\thesubsubsection}{0.5em}{}

\setlist[itemize]{topsep=3pt,itemsep=2pt,leftmargin=1.4em}
\setlist[enumerate]{topsep=3pt,itemsep=2pt,leftmargin=1.6em}
\setlength{\parskip}{0.45em}
\setlength{\parindent}{0pt}
\renewcommand{\arraystretch}{1.28}

\pagestyle{fancy}
\fancyhf{}
\fancyhead[L]{\small\color{SlateMid}PHOENIX evaluatiestudie -- Fase 1}
\fancyhead[R]{\small\color{SlateMid}\textbf{<<HCPCODE>>}\quad <<CA>> + <<CB>>}
\fancyfoot[C]{\small Pagina \thepage\ van \pageref{LastPage}}
\renewcommand{\headrulewidth}{0.3pt}
\renewcommand{\footrulewidth}{0pt}

\newtcolorbox{complaintbox}[1][]{
  colback=blue!3,colframe=PrimaryBlue,
  boxrule=0.8pt,arc=2.5mm,breakable,
  left=10pt,right=10pt,top=8pt,bottom=8pt,
  fonttitle=\small\bfseries\color{PrimaryBlue},
  title=Casusvignet,#1
}
\newtcolorbox{contextbox}[1][]{
  colback=green!3,colframe=ForestGreen,
  boxrule=0.6pt,arc=2mm,breakable,
  left=8pt,right=8pt,top=7pt,bottom=7pt,
  fonttitle=\small\bfseries\color{ForestGreen},
  title=Gestandaardiseerde context,#1
}
\newtcolorbox{instrbox}[1][]{
  colback=SoftBG,colframe=BorderGrey,
  boxrule=0.5pt,arc=1.5mm,breakable,
  left=8pt,right=8pt,top=7pt,bottom=7pt,
  fonttitle=\bfseries\color{SlateMid},#1
}
\newtcolorbox{monitorbox}{
  colback=SoftBG,colframe=BorderGrey,
  boxrule=0.5pt,arc=1.5mm,breakable,
  left=6pt,right=6pt,top=5pt,bottom=5pt
}
\newtcolorbox{responsebox}{
  colback=yellow!5,colframe=GoldAmber,
  boxrule=0.7pt,arc=2mm,breakable,
  left=10pt,right=10pt,top=9pt,bottom=9pt,
  fonttitle=\small\bfseries\color{GoldAmber},
  title=Uw antwoord
}

\tikzset{
  crnode/.style={
    draw=PrimaryBlue,fill=AccentBlue,rounded corners=2pt,
    text width=3.55cm,minimum height=0.76cm,
    font=\fontsize{6.5}{8}\selectfont\bfseries\color{DarkBlue},
    align=center,inner sep=3pt
  },
  prnode/.style={
    draw=ForestGreen,fill=AccentTeal,rounded corners=2pt,
    text width=3.55cm,minimum height=0.76cm,
    font=\fontsize{6.5}{8}\selectfont\bfseries\color{ForestGreen},
    align=center,inner sep=3pt
  },
  edgeS/.style={line width=1.6pt,draw=StrongRed!80,opacity=0.92},
  edgeM/.style={line width=1.0pt,draw=PrimaryBlue!70,opacity=0.88},
  edgeW/.style={line width=0.7pt,dashed,draw=gray!55,opacity=0.70},
}

\newcommand{\writeline}{\vspace{0.22em}\noindent\rule{\linewidth}{0.25pt}\vspace{0.42em}\par}

\newcommand{\criterionslot}[1]{%
\vspace{0.25em}\noindent
\begin{tcolorbox}[colback=SoftBG,colframe=BorderGrey,boxrule=0.4pt,arc=1.5mm,left=8pt,right=8pt,top=5pt,bottom=5pt]%
\small\textbf{\color{SlateMid}Criterium #1}\hspace{1em}%
\textbf{Label (2--5 woorden):}\hspace{0.5em}\rule{0.50\linewidth}{0.25pt}%
\end{tcolorbox}%
}

\newcommand{\predictorslot}[1]{%
\vspace{0.25em}\noindent
\begin{tcolorbox}[colback=SoftBG,colframe=BorderGrey,boxrule=0.4pt,arc=1.5mm,left=8pt,right=8pt,top=5pt,bottom=5pt]%
\small\textbf{\color{SlateMid}Predictor #1}\hspace{1em}%
\textbf{Label (2--6 woorden):}\hspace{0.5em}\rule{0.50\linewidth}{0.25pt}%
\end{tcolorbox}%
}

\newcommand{\rankbox}[1]{%
\vspace{0.22em}\noindent
\begin{tikzpicture}[baseline=3pt]
  \node[draw=DarkBlue,fill=AccentBlue,rounded corners=3pt,minimum width=2.5cm,minimum height=0.72cm,font=\small\bfseries\color{DarkBlue},align=center] {Prioriteit #1};
\end{tikzpicture}
\hspace{0.7em}\rule{0.60\linewidth}{0.25pt}\par\vspace{0.05em}
}

\newcommand{\checkitemBFS}[2]{%
\noindent\small\textbf{#1.}\hspace{0.45em}$\square$\hspace{0.4em}#2\par\vspace{0.20em}}

\newcommand{\messagelines}{\writeline\writeline\writeline\writeline}

\begin{document}
"""


def preamble(hcp_code: str, ca_label: str, cb_label: str) -> str:
    return (
        PREAMBLE_TEMPLATE.replace(
            "<<PDFTITLE>>",
            "PHOENIX evaluatiestudie - " + hcp_code + " (" + ca_label + " + " + cb_label + ")",
        )
        .replace("<<HCPCODE>>", hcp_code)
        .replace("<<CA>>", ca_label)
        .replace("<<CB>>", cb_label)
    )


COVER_TEMPLATE = r"""% ── TITELPAGINA ─────────────────────────────────────────────────────────────
\thispagestyle{fancy}
\vspace*{0.3cm}

\begin{tcolorbox}[
  colback=DarkBlue,colframe=DarkBlue,
  arc=3mm,left=14pt,right=14pt,top=12pt,bottom=12pt,
  width=\textwidth
]
\centering
{\fontsize{22}{28}\selectfont\bfseries\color{white}
PHOENIX evaluatiestudie\\[0.3em]
\fontsize{15}{20}\selectfont\color{AccentBlue}
Fase 1 -- onafhankelijke expertgeneratie\\[0.2em]
\fontsize{12}{16}\selectfont\color{white}
\textit{Instrument voor Zorgprofessionals}
}
\end{tcolorbox}

\vspace{0.6cm}

\begin{center}
\begin{tcolorbox}[
  colback=AccentBlue,colframe=PrimaryBlue,
  arc=2.5mm,boxrule=1.2pt,left=12pt,right=12pt,top=10pt,bottom=10pt,
  width=0.86\textwidth
]
\centering
{\large\bfseries\color{DarkBlue}Deelnemerscode: \texttt{<<HCPCODE>>}}\\[0.5em]
{\normalsize\color{SlateMid}
Toegewezen casussen:\quad
\textbf{\color{PrimaryBlue}<<CA>>} (<<CA_PROFILE>>)
\quad en \quad
\textbf{\color{PrimaryBlue}<<CB>>} (<<CB_PROFILE>>)
}
\end{tcolorbox}
\end{center}

\vspace{0.5cm}

\begin{center}
\begin{tcolorbox}[
  colback=SoftBG,colframe=BorderGrey,
  arc=2mm,left=10pt,right=10pt,top=9pt,bottom=9pt,
  width=0.92\textwidth
]
\small\renewcommand{\arraystretch}{1.28}
\begin{tabularx}{\textwidth}{@{}>{\bfseries\raggedright\arraybackslash}p{0.29\textwidth}X@{}}
Studie & Evaluatie van de klinische kwaliteit van een ontologiegebaseerd multi-agentsysteem voor gepersonaliseerde digitale geestelijke gezondheidszorg (PHOENIX) \\
Instelling & Universiteit Gent -- Faculteit Psychologie en Pedagogische Wetenschappen \\
Onderzoeker & Stijn Van Severen (masterproefstudent) \\
Promotoren & Prof.\ Dr.\ Geert Crombez; Dr.\ Annick De Paepe \\
Contact & \texttt{stijn.vanseveren@ugent.be} \\
Geschatte duur & Ongeveer 35--45 minuten voor beide casussen samen \\
\end{tabularx}
\end{tcolorbox}
\end{center}

\vspace{0.4cm}

\begin{center}
\begin{tcolorbox}[
  colback=SoftBG,colframe=BorderGrey,
  arc=2mm,left=10pt,right=10pt,top=9pt,bottom=9pt,
  width=0.92\textwidth
]
\textbf{\color{PrimaryBlue}Doel van deze bundel}

\medskip\small
U neemt deel aan een evaluatiestudie waarin zorgprofessionals onafhankelijk dezelfde
vijf klinische redeneerstappen uitvoeren als het PHOENIX-systeem. Uw antwoorden
vormen het menselijke referentiecorpus voor een latere dubbelblinde vergelijking
met systeemoutput.

\smallskip
Voor elk van uw twee casussen vult u dezelfde vijf delen in: (1) operationalisering,
(2) initieel observatiemodel, (3) prioritering van behandeldoelen, (4) verfijning van
EMA-metingen en (5) een mobiele coachingsboodschap.

\smallskip
\textbf{We vragen uw eigen klinische oordeelsvorming.} Antwoord zoals u dat in een
reele professionele context zou doen, maar werk strikt volgens de instructies op de
volgende pagina.
\end{tcolorbox}
\end{center}

\vspace{0.4cm}

\begin{center}
\begin{tcolorbox}[
  colback=AccentAmber,colframe=GoldAmber,
  arc=2mm,boxrule=0.7pt,left=10pt,right=10pt,top=8pt,bottom=8pt,
  width=0.90\textwidth
]
\small\centering
\textbf{\color{GoldAmber}Vertrouwelijkheid:} uw antwoorden worden voor analyse
geanonimiseerd en uitsluitend gebruikt binnen deze masterproefstudie. Deelname is
vrijwillig; u kan zich op elk moment terugtrekken door contact op te nemen met de
onderzoeker.
\end{tcolorbox}
\end{center}

\vfill
\newpage
\tableofcontents
\newpage
"""


def cover_page(hcp_code: str, ca: dict, cb: dict) -> str:
    return (
        COVER_TEMPLATE.replace("<<HCPCODE>>", hcp_code)
        .replace("<<CA>>", ca["label"])
        .replace("<<CB>>", cb["label"])
        .replace("<<CA_PROFILE>>", ca["profile"])
        .replace("<<CB_PROFILE>>", cb["profile"])
    )


INTRO = r"""
% ── INSTRUCTIEPAGINA ─────────────────────────────────────────────────────────
\section{Instructiepagina}

\textbf{PHOENIX} is een multi-agentsysteem dat een vrije klachttekst omzet in een
gestructureerde klinische redenering via vijf opeenvolgende stappen. In deze bundel
voert u diezelfde vijf stappen onafhankelijk uit voor uw twee toegewezen casussen.
Uw antwoorden vormen het menselijke referentiecorpus voor een latere dubbelblinde
vergelijking met systeemoutput.

\vspace{0.4em}
\begin{center}
\renewcommand{\arraystretch}{1.32}
\small
\begin{tabular}{@{}>{\bfseries\color{PrimaryBlue}}p{0.05\textwidth}>{\bfseries}p{0.28\textwidth}p{0.37\textwidth}p{0.15\textwidth}@{}}
\toprule
Stap & Klinische taak & Wat u doet & Richttijd \\
\midrule
1 & Operationalisering & Noteer 2--6 criteriumlabels voor actuele probleemdimensies & $\approx$\,6 min \\
2 & Initieel observatiemodel & Genereer 3--5 biopsychosociale predictorlabels (EMA-geschikt) & $\approx$\,6 min \\
3 & Behandeldoelprioritering & Rangschik de 5 standaardpredictoren van hoog naar laag & $\approx$\,7 min \\
4 & Verfijnd observatiemodel & Selecteer exact 5 EMA-items uit de lijst van 20 & $\approx$\,8 min \\
5 & Mobiele coaching & Schrijf een korte patientgerichte boodschap voor de app & $\approx$\,8 min \\
\midrule
 & \textbf{Totaal (2 casussen)} & & \textbf{35--45 min} \\
\bottomrule
\end{tabular}
\end{center}

\vspace{0.5em}
\begin{instrbox}[title={Werkwijze -- strikt te volgen}]
\begin{enumerate}[topsep=1pt,itemsep=2pt]
\item \textbf{Werk sequentieel: Deel 1 $\rightarrow$ Deel 5.}
Ga pas naar een volgend deel wanneer het huidige deel volledig is afgewerkt voor beide casussen.
\item \textbf{Gebruik geen generatieve AI, schrijfhulpmiddelen, richtlijnen of collegaoverleg.}
Extern gebruik ondermijnt de methodologische validiteit en de blind scoringswaarde van de studie.
\item \textbf{Gebruik in latere delen uitsluitend de meegeleverde gestandaardiseerde context.}
Die is bewust vastgezet zodat alle deelnemers op identieke input reageren.
\item \textbf{Herwerk eerdere antwoorden niet retroactief} nadat u latere context hebt gezien.
\item \textbf{Noteer of typ rechtstreeks in de voorziene antwoordzones.}
Onleesbare of ambigu geformuleerde antwoorden bemoeilijken latere blind beoordeling.
\end{enumerate}
\end{instrbox}

\vspace{0.4em}
\begin{instrbox}[title={EMA-principes (relevant voor Deel 2 en Deel 4)}]
\textbf{Ecological Momentary Assessment (EMA)} meet dagelijkse schommelingen via een mobiele
app. Elke EMA-variabele moet aan vier vereisten voldoen:
\begin{itemize}[topsep=2pt,itemsep=1pt]
\item \textbf{Dagelijks rapporteerbaar} via een korte vraag op de smartphone.
\item \textbf{Dynamisch en veranderbaar:} geen vaste diagnose, trait of achtergrondkenmerk.
\item \textbf{Meetbaar in een eenvoudig format:} ja/nee, aantal, minuten of een 0--10 score.
\item \textbf{Klinisch relevant voor opvolging:} toont binnen-persoonsvariatie die therapeutisch informatief is.
\end{itemize}
\smallskip
In \textbf{Deel 2} genereert u zelf predictorlabels. In \textbf{Deel 4} zijn de 20 kandidaat-items
reeds uitgewerkt als dagelijkse EMA-items; u selecteert de meest geschikte 5.
\end{instrbox}

\newpage
"""


PART1_HEADER = r"""
\section{Deel 1: Operationalisering van mentale gezondheidsproblemen}

\begin{tcolorbox}[colback=blue!3,colframe=PrimaryBlue,arc=2.5mm,boxrule=0.9pt,
  left=10pt,right=10pt,top=9pt,bottom=9pt]
\textbf{\color{PrimaryBlue}Instructies Deel 1}

\smallskip\small
\textbf{Opdracht:} identificeer de belangrijkste actuele probleemdimensies in de klachttekst en noteer
voor elke dimensie uitsluitend een \textbf{kort criteriumlabel}.

\textbf{Wat telt als criterium?}
Een klinisch relevante probleemdimensie die:
\begin{itemize}[topsep=2pt,itemsep=1pt]
\item momenteel aanwezig is in de casus,
\item inhoudelijk apart te onderscheiden is van andere probleemdimensies,
\item in principe herhaald meetbaar zou kunnen zijn.
\end{itemize}

\textbf{Antwoordformat:} label van 2--5 woorden, zonder beschrijvende zin.
Noteer per casus \textbf{2--6 criteria}. Laat ongebruikte velden leeg.
\end{tcolorbox}

\vspace{0.5em}
"""


PART1_CASE_TEMPLATE = r"""
\subsection*{Casus <<LABEL>>: <<PROFILE>>}

\begin{complaintbox}[title={Casus <<LABEL>>: <<PROFILE>>; duur: <<DURATION>>}]
\small <<VIGNETTE>>
\end{complaintbox}

\begin{responsebox}
\small\textbf{Opdracht:} noteer \textbf{2--6 criteriumlabels} voor deze casus.
Gebruik enkel korte labels; voeg geen beschrijving toe.

\criterionslot{1}
\criterionslot{2}
\criterionslot{3}
\criterionslot{4}
\criterionslot{5}
\criterionslot{6}
\end{responsebox}

\vspace{0.6em}
"""


def part1_case(case: dict) -> str:
    return (
        PART1_CASE_TEMPLATE.replace("<<LABEL>>", case["label"])
        .replace("<<PROFILE>>", case["profile"])
        .replace("<<DURATION>>", case["duration"])
        .replace("<<VIGNETTE>>", case["vignette"])
    )


def part1(case_a: dict, case_b: dict) -> str:
    return PART1_HEADER + "\n\\newpage\n" + part1_case(case_a) + "\n\\newpage\n" + part1_case(case_b) + "\n\\newpage\n"


PART2_HEADER = r"""
\section{Deel 2: Initieel observatiemodel}

\begin{tcolorbox}[colback=green!3,colframe=ForestGreen,arc=2.5mm,boxrule=0.9pt,
  left=10pt,right=10pt,top=9pt,bottom=9pt]
\textbf{\color{ForestGreen}Instructies Deel 2}

\smallskip\small
\textbf{Opdracht:} genereer \textbf{3--5 biopsychosociale predictorlabels} die een initieel observatiemodel
vormen voor deze casus.

\textbf{Wat telt als predictor?}
Een veranderbare factor die:
\begin{itemize}[topsep=2pt,itemsep=1pt]
\item klinisch plausibel samenhangt met een of meerdere criteria,
\item geschikt is voor \textbf{dagelijkse Ecological Momentary Assessment (EMA)},
\item door de persoon zelf eenvoudig via een mobiele app kan worden gerapporteerd.
\end{itemize}

\textbf{Antwoordformat:} enkel een label van 2--6 woorden, zonder meetdefinitie of toelichting.
Laat ongebruikte velden leeg.
\end{tcolorbox}

\vspace{0.5em}
"""


def part2_case(case: dict) -> str:
    criteria_items = "\n".join(
        rf"\item \textbf{{CR-{idx}}} {label}" for idx, label in enumerate(case["criteria"], start=1)
    )
    return J(
        rf"\subsection*{{Casus {case['label']}: Deel 2}}",
        "",
        rf"\begin{{complaintbox}}[title={{Casus {case['label']}: verkorte klachtomschrijving}}]",
        rf"\small {case['profile']}; duur: {case['duration']}. {case['short_desc']}",
        r"\end{complaintbox}",
        "",
        r"\begin{contextbox}",
        r"\small\textbf{Gestandaardiseerde criteria uit stap 1:}",
        r"\begin{itemize}[topsep=2pt,itemsep=1pt]",
        criteria_items,
        r"\end{itemize}",
        r"\end{contextbox}",
        "",
        r"\begin{responsebox}",
        r"\small\textbf{Opdracht:} noteer \textbf{3--5 predictorlabels} voor een initieel observatiemodel.",
        r"Zorg dat elk label later dagelijks via een mobiele app meetbaar kan worden gemaakt.",
        "",
        r"\predictorslot{1}",
        r"\predictorslot{2}",
        r"\predictorslot{3}",
        r"\predictorslot{4}",
        r"\predictorslot{5}",
        r"\end{responsebox}",
        "",
        r"\vspace{0.6em}",
    )


def part2(case_a: dict, case_b: dict) -> str:
    return PART2_HEADER + "\n\\newpage\n" + part2_case(case_a) + "\n\\newpage\n" + part2_case(case_b) + "\n\\newpage\n"


PART3_HEADER = r"""
\section{Deel 3: Prioritering van behandeldoelen}

\begin{tcolorbox}[colback=yellow!5,colframe=GoldAmber,arc=2.5mm,boxrule=0.9pt,
  left=10pt,right=10pt,top=9pt,bottom=9pt]
\textbf{\color{GoldAmber}Instructies Deel 3}

\smallskip\small
\textbf{Opdracht:} rangschik de \textbf{5 standaardpredictoren} van \textbf{hoogste} naar \textbf{laagste}
klinische prioriteit als behandeldoel.

\textbf{Gebruik voor uw rangschikking:}
\begin{itemize}[topsep=2pt,itemsep=1pt]
\item de 21-daagse monitoring,
\item de bipartiete netwerkstructuur,
\item de mate waarin een predictor meerdere criteria kan beinvloeden.
\end{itemize}

\textbf{Visuele sleutel van het netwerk:}
\begin{center}
\begin{tikzpicture}[baseline=-0.5ex]
\draw[line width=1.6pt,draw=StrongRed!80] (0,0) -- (1.0,0);
\end{tikzpicture}~\textbf{Rood = sterke relatie}\qquad
\begin{tikzpicture}[baseline=-0.5ex]
\draw[line width=1.0pt,draw=PrimaryBlue!70] (0,0) -- (1.0,0);
\end{tikzpicture}~\textbf{Blauw = matige relatie}
\end{center}

\textbf{Antwoordformat:} vul alle prioriteitslijnen in, van 1 tot en met 5.
\end{tcolorbox}

\vspace{0.5em}
"""


CR_Y = [1.65, 0.55, -0.55, -1.65]
PR_Y = [2.20, 1.10, 0.00, -1.10, -2.20]


def tikz_network(case: dict) -> str:
    criteria_nodes = "\n  ".join(
        rf"\node[crnode] (cr{idx}) at (0, {CR_Y[idx-1]}) {{CR-{idx}\\{label}}};"
        for idx, label in enumerate(case["criteria"], start=1)
    )
    predictor_nodes = "\n  ".join(
        rf"\node[prnode] (p{idx}) at (9, {PR_Y[idx-1]}) {{P{idx}\\{label}}};"
        for idx, label in enumerate(case["predictors"], start=1)
    )
    edge_map = {"S": "edgeS", "M": "edgeM", "W": "edgeW"}
    edges = "\n  ".join(
        rf"\draw[{edge_map[weight]}] ({src}.east) -- ({dst}.west);"
        for src, dst, weight in case["tikz_edges"]
    )
    return J(
        r"\begin{center}",
        r"\begin{tikzpicture}[node distance=0pt]",
        "  " + criteria_nodes,
        "  " + predictor_nodes,
        "  " + edges,
        r"\end{tikzpicture}",
        r"\end{center}",
    )


def part3_case(case: dict) -> str:
    predictor_line = r",\quad ".join(
        rf"\textbf{{P{idx}: {label}}}" for idx, label in enumerate(case["predictors"], start=1)
    )
    rank_boxes = "\n".join(rf"\rankbox{{{idx}}}" for idx in range(1, len(case["predictors"]) + 1))
    return J(
        rf"\subsection*{{Casus {case['label']}: Deel 3}}",
        "",
        rf"\begin{{complaintbox}}[title={{Casus {case['label']}: {case['profile']}}}]",
        rf"\small {case['short_desc']} Duur: {case['duration']}.",
        r"\end{complaintbox}",
        "",
        r"\begin{monitorbox}",
        rf"\small\textbf{{21-daagse monitoring:}} {case['monitoring']}",
        r"\end{monitorbox}",
        "",
        tikz_network(case),
        "",
        r"\begin{responsebox}",
        r"\small\textbf{Opdracht:} rangschik \textbf{alle 5 predictors} van hoogste naar laagste behandelprioriteit.",
        rf"\textbf{{Beschikbare predictors:}} {predictor_line}",
        "",
        rank_boxes,
        r"\end{responsebox}",
        "",
        r"\vspace{0.6em}",
    )


def part3(case_a: dict, case_b: dict) -> str:
    return PART3_HEADER + "\n\\newpage\n" + part3_case(case_a) + "\n\\newpage\n" + part3_case(case_b) + "\n\\newpage\n"


PART4_HEADER = r"""
\section{Deel 4: Verfijnd observatiemodel via breadth-first update-logica}

\begin{tcolorbox}[colback=purple!4,colframe=RichPurple,arc=2.5mm,boxrule=0.9pt,
  left=10pt,right=10pt,top=9pt,bottom=9pt]
\textbf{\color{RichPurple}Instructies Deel 4}

\smallskip\small
\textbf{Opdracht:} selecteer per casus \textbf{exact 5 EMA-items} uit een lijst van 20.

\textbf{Wat bootst dit deel na?}
PHOENIX verfijnt het observatiemodel via een \textbf{breadth-first update-logica}: vanuit de
reeds geprioriteerde behandeldoelen wordt gezocht naar de meest geschikte, direct meetbare
subpredictoren voor de volgende EMA-cyclus.

\textbf{Selectieprincipe:}
\begin{itemize}[topsep=2pt,itemsep=1pt]
\item start vanuit de standaardbehandeldoelen hieronder,
\item kies de 5 dagelijkse EMA-items die daar het best op aansluiten,
\item verkies klinisch relevante breedte boven irrelevante of perifere items,
\item noteer \textbf{geen} toelichting of extra commentaar.
\end{itemize}

\textbf{Antwoordformat:} vink \textbf{exact 5} items aan.
\end{tcolorbox}

\vspace{0.5em}
"""


def bfs_items_block(items: list[str]) -> str:
    return "\n".join(rf"\checkitemBFS{{{idx}}}{{{label}}}" for idx, label in enumerate(items, start=1))


def part4_case(case: dict) -> str:
    targets = "\n".join(
        rf"\item \textbf{{{label}}}\quad\textit{{\small ({rationale})}}"
        for label, rationale in case["treatment_targets"]
    )
    return J(
        rf"\subsection*{{Casus {case['label']}: Deel 4}}",
        "",
        r"\begin{contextbox}",
        r"\small\textbf{Gestandaardiseerde behandeldoelen uit Deel 3:}",
        r"\begin{enumerate}[topsep=2pt,itemsep=1pt,label=\arabic*.]",
        targets,
        r"\end{enumerate}",
        r"\end{contextbox}",
        "",
        r"\begin{responsebox}",
        r"\small\textbf{Opdracht:} selecteer \textbf{exact 5} EMA-items die het best passen als volgende subpredictoren.",
        r"\textit{Alle antwoordopties zijn dagelijkse mobiele EMA-items.}",
        r"",
        bfs_items_block(case["bfs_items"]),
        r"\vspace{0.4em}",
        r"\noindent\textbf{Totaal geselecteerd:}\hspace{0.5em}\rule{1.2cm}{0.25pt}\hspace{0.2em}/ 5",
        r"\end{responsebox}",
        "",
        r"\vspace{0.6em}",
    )


def part4(case_a: dict, case_b: dict) -> str:
    return PART4_HEADER + "\n\\newpage\n" + part4_case(case_a) + "\n\\newpage\n" + part4_case(case_b) + "\n\\newpage\n"


PART5_HEADER = r"""
\section{Deel 5: Mobiele coachingsboodschap}

\begin{tcolorbox}[colback=red!3,colframe=ForestGreen,arc=2.5mm,boxrule=0.9pt,
  left=10pt,right=10pt,top=9pt,bottom=9pt]
\textbf{\color{ForestGreen}Instructies Deel 5}

\smallskip\small
\textbf{Opdracht:} schrijf een korte, patientgerichte coachingsboodschap die rechtstreeks in de
mobiele applicatie kan verschijnen.

\textbf{Gebruik hiervoor:}
\begin{itemize}[topsep=2pt,itemsep=2pt]
\item het primaire probleem van de casus,
\item het geselecteerde behandeldoel,
\item de voornaamste barriere,
\item de aangegeven copingstrategie.
\end{itemize}

\textbf{De boodschap hoort:}
\begin{itemize}[topsep=2pt,itemsep=2pt]
\item compact genoeg te zijn voor een mobiel scherm,
\item warm en professioneel te klinken,
\item een concrete eerstvolgende stap te bevatten,
\item rechtstreeks tot de persoon gericht te zijn.
\end{itemize}

\textbf{Werkvoorbeeld} (niet gerelateerd aan een studiecasus):
\begin{tcolorbox}[colback=SoftBG,colframe=BorderGrey,boxrule=0.4pt,arc=1.5mm,
  left=7pt,right=7pt,top=5pt,bottom=5pt]
\small
\textbf{Context:} verpleegkundige met het doel om opnieuw korte beweegmomenten in te bouwen;
barriere = vermoeidheid na de shift; coping = een vaste 10-minutenwandeling koppelen aan het thuiskomen.\\[0.3em]
\textbf{Voorbeeldboodschap:}\\
\textit{``Na een zware shift voelt rust nemen logisch, maar net dat eerste kleine beweegmoment kan je avond helpen ontladen. Trek vanavond meteen na thuiskomst je schoenen aan en wandel 10 minuten buiten, zonder jezelf meer op te leggen dan dat ene blokje. Zo maak je de stap haalbaar en vergroot je de kans dat je lichaam later echt kan afschakelen.''}
\end{tcolorbox}
\end{tcolorbox}

\vspace{0.5em}
"""


def part5_case(case: dict) -> str:
    return J(
        rf"\subsection*{{Casus {case['label']}: Deel 5}}",
        "",
        r"\begin{contextbox}",
        r"\small\renewcommand{\arraystretch}{1.24}",
        r"\begin{tabular}[t]{@{}>{\raggedright\bfseries\color{ForestGreen}\arraybackslash}p{0.27\linewidth}>{\raggedright\arraybackslash}p{0.67\linewidth}@{}}",
        rf"Primair probleem & {case['p5_challenge']} \\[0.15em]",
        rf"Behandeldoel & {case['p5_target']} \\[0.15em]",
        rf"Voornaamste barriere & {case['p5_barrier']} \\[0.15em]",
        rf"Copingstrategie & {case['p5_coping']} \\",
        r"\end{tabular}",
        r"\end{contextbox}",
        "",
        r"\begin{responsebox}",
        r"\small\textbf{Opdracht:} schrijf hieronder de mobiele coachingsboodschap voor deze casus.",
        r"\textit{Formuleer alsof de tekst morgen rechtstreeks op de smartphone van de persoon verschijnt.}",
        r"",
        r"\messagelines",
        r"\end{responsebox}",
        "",
        r"\vspace{0.6em}",
    )


def part5(case_a: dict, case_b: dict) -> str:
    return PART5_HEADER + "\n\\newpage\n" + part5_case(case_a) + "\n\\newpage\n" + part5_case(case_b) + "\n\\newpage\n"


COMPLETION_TEMPLATE = r"""
\section{Afronding en terugbezorging}

Dank u voor het invullen van alle vijf delen voor uw twee toegewezen casussen
(<<CA>> en <<CB>>).

\begin{instrbox}[title=Checklist voor terugbezorging]
\begin{itemize}
\item[$\square$] Ik heb in Deel 1 voor beide casussen criteriumlabels ingevuld.
\item[$\square$] Ik heb in Deel 2 voor beide casussen predictorlabels ingevuld.
\item[$\square$] Ik heb in Deel 3 voor beide casussen alle 5 predictors gerangschikt.
\item[$\square$] Ik heb in Deel 4 voor beide casussen exact 5 EMA-items geselecteerd.
\item[$\square$] Ik heb in Deel 5 voor beide casussen een mobiele coachingsboodschap geschreven.
\item[$\square$] Mijn antwoorden weerspiegelen mijn eigen klinische oordeel zonder gebruik van generatieve AI of andere externe hulp.
\item[$\square$] Ik begrijp dat mijn antwoorden geanonimiseerd worden voor analyse.
\end{itemize}
\end{instrbox}

\medskip
Bezorg het ingevulde document terug via e-mail aan:
\begin{center}
\texttt{stijn.vanseveren@ugent.be}\quad met onderwerp:\quad \texttt{PHOENIX-PRE-<<HCPCODE>>}
\end{center}

\vspace{0.5em}
\begin{tcolorbox}[colback=AccentBlue,colframe=PrimaryBlue,arc=2mm,boxrule=0.6pt,left=8pt,right=8pt,top=7pt,bottom=7pt]
\small\centering
Na ontvangst worden uw antwoorden geanonimiseerd en opgenomen in het expertreferentiecorpus
voor de latere blind evaluatie van PHOENIX. Hartelijk dank voor uw bijdrage aan dit onderzoek.
\end{tcolorbox}

\end{document}
"""


def completion_page(hcp_code: str, case_a: dict, case_b: dict) -> str:
    return (
        COMPLETION_TEMPLATE.replace("<<HCPCODE>>", hcp_code)
        .replace("<<CA>>", case_a["label"])
        .replace("<<CB>>", case_b["label"])
    )


def build_document(hcp_num: int) -> str:
    hcp_code, case_id_a, case_id_b = ASSIGNMENT[hcp_num]
    case_a = CASES[case_id_a]
    case_b = CASES[case_id_b]
    return J(
        preamble(hcp_code, case_id_a, case_id_b),
        cover_page(hcp_code, case_a, case_b),
        INTRO,
        part1(case_a, case_b),
        part2(case_a, case_b),
        part3(case_a, case_b),
        part4(case_a, case_b),
        part5(case_a, case_b),
        completion_page(hcp_code, case_a, case_b),
    )


BASE = Path(__file__).parent


def write_and_compile(hcp_num: int) -> None:
    out_dir = BASE / f"HCP_{hcp_num}"
    out_dir.mkdir(parents=True, exist_ok=True)
    tex_path = out_dir / "main.tex"

    tex_path.write_text(build_document(hcp_num), encoding="utf-8")
    print(f"[+] Geschreven: {tex_path}")

    result = subprocess.run(
        ["tectonic", str(tex_path)],
        capture_output=True,
        text=True,
        cwd=str(out_dir),
    )
    if result.returncode == 0:
        print(f"[v] Gecompileerd: HCP_{hcp_num}/main.pdf")
        return

    print(f"[X] Compilatiefout voor HCP_{hcp_num}")
    print(result.stdout[-3000:])
    print(result.stderr[-3000:])
    raise SystemExit(result.returncode)


# ── WORD-DOCUMENT GENERATIE ────────────────────────────────────────────────


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _shade_para(p, shade_hex: str) -> None:
    """Apply background shade to an existing paragraph."""
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), shade_hex)
    pPr.append(shd)


def _shade_cell(cell, shade_hex: str) -> None:
    """Apply background shade to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), shade_hex)
    tcPr.append(shd)


def _add_shaded_para(doc: Document, text: str, shade_hex: str = "EFF6FF",
                     align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
                     bold_first_line: bool = False) -> None:
    """Add a shaded paragraph. If bold_first_line, the first line is bold."""
    p = doc.add_paragraph()
    p.alignment = align
    _shade_para(p, shade_hex)
    lines = text.split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.size = Pt(10)
        if bold_first_line and i == 0:
            run.bold = True


def _add_centered_box(doc: Document, text: str, shade_hex: str,
                      bold_first_line: bool = False) -> None:
    """Add a full-width shaded box with centered text."""
    _add_shaded_para(doc, text, shade_hex=shade_hex,
                     align=WD_ALIGN_PARAGRAPH.CENTER,
                     bold_first_line=bold_first_line)


def _add_fillline(doc: Document, label: str, width_chars: int = 60) -> None:
    p = doc.add_paragraph()
    run_label = p.add_run(label + "  ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_line = p.add_run("_" * width_chars)
    run_line.font.size = Pt(10)
    run_line.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


def _add_checkbox_item(doc: Document, num: int, text: str) -> None:
    p = doc.add_paragraph(style="List Paragraph")
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f"{num:2d}.  \u25a1  {text}")
    run.font.size = Pt(10)


def _page_break(doc: Document) -> None:
    doc.add_page_break()


def _centered_meta_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    """Render a 2-column key/value table, centered on the page."""
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    col_widths = [Cm(4.2), Cm(10.5)]
    for row_idx, (k, v) in enumerate(rows):
        cells = tbl.rows[row_idx].cells
        cells[0].width = col_widths[0]
        cells[1].width = col_widths[1]
        _shade_cell(cells[0], "DBEAFE")
        _shade_cell(cells[1], "F8FAFC")
        rk = cells[0].paragraphs[0].add_run(k)
        rk.bold = True
        rk.font.size = Pt(9)
        rv = cells[1].paragraphs[0].add_run(v)
        rv.font.size = Pt(9)


def build_word_document(hcp_num: int) -> Document:
    hcp_code, case_id_a, case_id_b = ASSIGNMENT[hcp_num]
    case_a = CASES[case_id_a]
    case_b = CASES[case_id_b]

    doc = Document()

    # ── Pagina-marges ─────────────────────────────────────────────────────
    for sec in doc.sections:
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)

    # ── TITELPAGINA ────────────────────────────────────────────────────────
    # --- Hoofdtitel (gecentreerd, donkerblauw) ---
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("PHOENIX evaluatiestudie -- Fase 1")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    t.add_run("\n")
    r2 = t.add_run("Instrument voor Zorgprofessionals")
    r2.bold = True
    r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)

    doc.add_paragraph()

    # --- Deelnemerscode + casussen (gecentreerd, blauw) ---
    _add_centered_box(
        doc,
        f"Deelnemerscode:  {hcp_code}\n\n"
        f"Toegewezen casussen:\n"
        f"{case_a['label']}  ({case_a['profile']})\n"
        f"en\n"
        f"{case_b['label']}  ({case_b['profile']})",
        shade_hex="DBEAFE",
        bold_first_line=True,
    )

    doc.add_paragraph()

    # --- Studieinfo-tabel (gecentreerd) ---
    meta_lines = [
        ("Studie",
         "Evaluatie van de klinische kwaliteit van een ontologiegebaseerd "
         "multi-agentsysteem voor gepersonaliseerde digitale geestelijke "
         "gezondheidszorg (PHOENIX)"),
        ("Instelling",
         "Universiteit Gent -- Faculteit Psychologie en Pedagogische Wetenschappen"),
        ("Onderzoeker", "Stijn Van Severen (masterproefstudent)"),
        ("Promotoren",  "Prof. Dr. Geert Crombez; Dr. Annick De Paepe"),
        ("Contact",     "stijn.vanseveren@ugent.be"),
        ("Geschatte duur", "Ongeveer 35-45 minuten voor beide casussen samen"),
    ]
    _centered_meta_table(doc, meta_lines)

    doc.add_paragraph()

    # --- Doel van deze bundel ---
    _add_shaded_para(
        doc,
        "Doel van deze bundel\n\n"
        "U neemt deel aan een evaluatiestudie waarin zorgprofessionals onafhankelijk "
        "dezelfde vijf klinische redeneerstappen uitvoeren als het PHOENIX-systeem. "
        "Uw antwoorden vormen het menselijke referentiecorpus voor een latere "
        "dubbelblinde vergelijking met systeemoutput.\n\n"
        "Voor elk van uw twee casussen vult u dezelfde vijf delen in: "
        "(1) operationalisering, (2) initieel observatiemodel, "
        "(3) prioritering van behandeldoelen, (4) verfijning van EMA-metingen "
        "en (5) een mobiele coachingsboodschap.\n\n"
        "We vragen uw eigen klinische oordeelsvorming. Antwoord zoals u dat in een "
        "reele professionele context zou doen, maar werk strikt volgens de instructies "
        "op de volgende pagina.",
        shade_hex="EFF6FF",
        bold_first_line=True,
    )

    doc.add_paragraph()

    # --- Vertrouwelijkheid & geïnformeerde toestemming ---
    _add_shaded_para(
        doc,
        "Vertrouwelijkheid en geïnformeerde toestemming\n\n"
        "Uw antwoorden worden voor analyse geanonimiseerd en uitsluitend gebruikt "
        "binnen deze masterproefstudie. Na ontvangst worden ze opgenomen in het "
        "expertreferentiecorpus voor de latere blind evaluatie van het PHOENIX-systeem. "
        "Uw gegevens worden op geen enkel moment gekoppeld aan uw naam of identiteit "
        "in publieke of wetenschappelijke rapportage.\n\n"
        "Deelname is volledig vrijwillig. U kan zich op elk moment zonder opgave van "
        "reden terugtrekken door contact op te nemen met de onderzoeker via "
        "stijn.vanseveren@ugent.be. Terugtrekking heeft geen gevolgen.\n\n"
        "Door dit ingevulde document te bezorgen aan de onderzoeker bevestigt u dat:\n"
        "  (1)  u de bovenstaande informatie hebt gelezen en begrepen;\n"
        "  (2)  u vrijwillig instemt met deelname aan deze studie;\n"
        "  (3)  u begrijpt dat uw antwoorden geanonimiseerd worden verwerkt.\n\n"
        "Uw inzending geldt als geïnformeerde toestemming.",
        shade_hex="FEF3C7",
        bold_first_line=True,
    )

    _page_break(doc)

    # ── INHOUDSOPGAVE ─────────────────────────────────────────────────────
    _add_heading(doc, "Inhoudsopgave", level=1)
    toc_entries = [
        "1  Instructiepagina",
        "2  Deel 1: Operationalisering van mentale gezondheidsproblemen",
        "3  Deel 2: Initieel observatiemodel",
        "4  Deel 3: Prioritering van behandeldoelen",
        "5  Deel 4: Verfijnd observatiemodel via breadth-first update-logica",
        "6  Deel 5: Mobiele coachingsboodschap",
        "7  Afronding en terugbezorging",
    ]
    for entry in toc_entries:
        p = doc.add_paragraph(entry)
        p.paragraph_format.left_indent = Cm(0.5)
        for run in p.runs:
            run.font.size = Pt(10)

    _page_break(doc)

    # ── SECTIE 1: INSTRUCTIEPAGINA ─────────────────────────────────────────
    _add_heading(doc, "1  Instructiepagina", level=1)

    p = doc.add_paragraph(
        "PHOENIX is een multi-agentsysteem dat een vrije klachttekst omzet in een "
        "gestructureerde klinische redenering via vijf opeenvolgende stappen. In deze "
        "bundel voert u diezelfde vijf stappen onafhankelijk uit voor uw twee toegewezen "
        "casussen. Uw antwoorden vormen het menselijke referentiecorpus voor een latere "
        "dubbelblinde vergelijking met systeemoutput."
    )
    p.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # 5-stappen tabel
    steps = [
        ("1", "Operationalisering", "Noteer 2-6 criteriumlabels voor actuele probleemdimensies", "~6 min"),
        ("2", "Initieel observatiemodel", "Genereer 3-5 biopsychosociale predictorlabels (EMA-geschikt)", "~6 min"),
        ("3", "Behandeldoelprioritering", "Rangschik de 5 standaardpredictoren van hoog naar laag", "~7 min"),
        ("4", "Verfijnd observatiemodel", "Selecteer exact 5 EMA-items uit de lijst van 20", "~8 min"),
        ("5", "Mobiele coaching", "Schrijf een korte patientgerichte boodschap voor de app", "~8 min"),
    ]
    step_tbl = doc.add_table(rows=len(steps) + 1, cols=4)
    step_tbl.style = "Table Grid"
    headers = ["Stap", "Klinische taak", "Wat u doet", "Richttijd"]
    for col_i, hdr in enumerate(headers):
        cell = step_tbl.rows[0].cells[col_i]
        run = cell.paragraphs[0].add_run(hdr)
        run.bold = True
        run.font.size = Pt(9)
    for row_i, (stap, taak, doen, tijd) in enumerate(steps, start=1):
        cells = step_tbl.rows[row_i].cells
        for col_i, val in enumerate([stap, taak, doen, tijd]):
            run = cells[col_i].paragraphs[0].add_run(val)
            run.font.size = Pt(9)
            if col_i == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)

    doc.add_paragraph()
    _add_shaded_para(
        doc,
        "WERKWIJZE -- STRIKT TE VOLGEN:\n"
        "1. Werk sequentieel: Deel 1 -> Deel 5. Ga pas naar een volgend deel wanneer het huidige "
        "volledig is afgewerkt voor beide casussen.\n"
        "2. Gebruik geen generatieve AI, schrijfhulpmiddelen, richtlijnen of collegaoverleg. "
        "Extern gebruik ondermijnt de methodologische validiteit en de blind scoringswaarde van de studie.\n"
        "3. Gebruik in latere delen uitsluitend de meegeleverde gestandaardiseerde context. "
        "Die is bewust vastgezet zodat alle deelnemers op identieke input reageren.\n"
        "4. Herwerk eerdere antwoorden NIET retroactief nadat u latere context hebt gezien.\n"
        "5. Noteer of typ rechtstreeks in de voorziene antwoordzones. "
        "Onleesbare of ambigu geformuleerde antwoorden bemoeilijken latere blind beoordeling.",
        shade_hex="FFF7ED",
    )

    doc.add_paragraph()
    _add_shaded_para(
        doc,
        "EMA-PRINCIPES (relevant voor Deel 2 en Deel 4):\n"
        "Ecological Momentary Assessment (EMA) meet dagelijkse schommelingen via een mobiele app. "
        "Elke EMA-variabele moet aan vier vereisten voldoen:\n"
        "  (1) Dagelijks rapporteerbaar via een korte smartphone-vraag.\n"
        "  (2) Dynamisch en veranderbaar -- geen vaste diagnose, trait of achtergrondkenmerk.\n"
        "  (3) Meetbaar in een eenvoudig format: ja/nee, aantal, minuten of een 0-10 score.\n"
        "  (4) Klinisch relevant voor opvolging -- toont binnen-persoonsvariatie die therapeutisch informatief is.\n\n"
        "In Deel 2 genereert u zelf predictorlabels. In Deel 4 zijn de 20 kandidaat-items reeds "
        "uitgewerkt als dagelijkse EMA-items; u selecteert de meest geschikte 5.",
        shade_hex="ECFDF5",
    )

    _page_break(doc)

    # ── DEEL 1 ─────────────────────────────────────────────────────────────
    def word_part1_case(case: dict) -> None:
        _add_heading(doc, f"Casus {case['label']}: {case['profile']}", level=2)
        _add_shaded_para(
            doc,
            f"CASUSVIGNET  |  Duur: {case['duration']}\n\n{case['vignette']}",
            shade_hex="EFF6FF",
        )
        doc.add_paragraph()
        instr = doc.add_paragraph()
        ri2 = instr.add_run(
            "Opdracht: noteer 2-6 criteriumlabels voor deze casus. "
            "Gebruik enkel korte labels (2-5 woorden); voeg geen beschrijving toe."
        )
        ri2.font.size = Pt(10)
        ri2.bold = True
        for i in range(1, 7):
            _add_fillline(doc, f"Criterium {i}  Label (2-5 woorden):", 50)

    _add_heading(doc, "2  Deel 1: Operationalisering van mentale gezondheidsproblemen", level=1)
    _add_shaded_para(
        doc,
        "Identificeer de belangrijkste actuele probleemdimensies in de klachttekst en noteer voor "
        "elke dimensie uitsluitend een kort criteriumlabel (2-5 woorden). "
        "Noteer 2-6 criteria per casus; laat ongebruikte velden leeg.",
        shade_hex="DBEAFE",
    )
    _page_break(doc)
    word_part1_case(case_a)
    _page_break(doc)
    word_part1_case(case_b)
    _page_break(doc)

    # ── DEEL 2 ─────────────────────────────────────────────────────────────
    def word_part2_case(case: dict) -> None:
        _add_heading(doc, f"Casus {case['label']}: Deel 2", level=2)
        _add_shaded_para(
            doc,
            f"Casus {case['label']}: verkorte klachtomschrijving\n"
            f"{case['profile']}; duur: {case['duration']}. {case['short_desc']}",
            shade_hex="EFF6FF",
        )
        doc.add_paragraph()
        ctx = doc.add_paragraph()
        ctx.add_run("Gestandaardiseerde criteria uit stap 1:").bold = True
        ctx.runs[0].font.size = Pt(10)
        for idx, lbl in enumerate(case["criteria"], start=1):
            p = doc.add_paragraph(f"  CR-{idx}  {lbl}", style="List Bullet")
            p.runs[0].font.size = Pt(10)
        doc.add_paragraph()
        instr = doc.add_paragraph()
        ri3 = instr.add_run(
            "Opdracht: noteer 3-5 predictorlabels (2-6 woorden) voor een initieel observatiemodel. "
            "Elk label moet later dagelijks via een mobiele app meetbaar kunnen worden gemaakt."
        )
        ri3.bold = True
        ri3.font.size = Pt(10)
        for i in range(1, 6):
            _add_fillline(doc, f"Predictor {i}  Label (2-6 woorden):", 50)

    _add_heading(doc, "3  Deel 2: Initieel observatiemodel", level=1)
    _add_shaded_para(
        doc,
        "Genereer 3-5 biopsychosociale predictorlabels die een initieel observatiemodel vormen. "
        "Elke predictor moet klinisch plausibel samenhangen met een of meerdere criteria en "
        "geschikt zijn voor dagelijkse EMA (zie instructiepagina).",
        shade_hex="D1FAE5",
    )
    _page_break(doc)
    word_part2_case(case_a)
    _page_break(doc)
    word_part2_case(case_b)
    _page_break(doc)

    # ── DEEL 3 ─────────────────────────────────────────────────────────────
    def word_part3_case(case: dict) -> None:
        _add_heading(doc, f"Casus {case['label']}: Deel 3", level=2)
        _add_shaded_para(
            doc,
            f"Casus {case['label']}: {case['profile']}\n{case['short_desc']}  Duur: {case['duration']}.",
            shade_hex="EFF6FF",
        )
        mon = doc.add_paragraph()
        mon.add_run("21-daagse monitoring:  ").bold = True
        mon.add_run(case["monitoring"]).font.size = Pt(10)
        mon.runs[0].font.size = Pt(10)
        doc.add_paragraph()

        # Network beschrijving als tekst
        net = doc.add_paragraph()
        net.add_run("Bipartiet netwerk (zie PDF voor visuele weergave):").bold = True
        net.runs[0].font.size = Pt(10)

        strong_edges = [(s, d) for s, d, w in case["tikz_edges"] if w == "S"]
        mod_edges = [(s, d) for s, d, w in case["tikz_edges"] if w == "M"]

        def edge_label(node: str) -> str:
            if node.startswith("cr"):
                idx = int(node[2:])
                return f"CR-{idx} ({case['criteria'][idx-1]})"
            else:
                idx = int(node[1:])
                return f"P{idx} ({case['predictors'][idx-1]})"

        if strong_edges:
            p = doc.add_paragraph()
            p.add_run("  Sterke relaties (rood):  ").bold = True
            p.runs[0].font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)
            p.runs[0].font.size = Pt(9)
            p.add_run(",  ".join(f"{edge_label(s)} <-> {edge_label(d)}" for s, d in strong_edges)).font.size = Pt(9)
        if mod_edges:
            p = doc.add_paragraph()
            p.add_run("  Matige relaties (blauw):  ").bold = True
            p.runs[0].font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
            p.runs[0].font.size = Pt(9)
            p.add_run(",  ".join(f"{edge_label(s)} <-> {edge_label(d)}" for s, d in mod_edges)).font.size = Pt(9)

        predictor_list = ",  ".join(
            f"P{i}: {lbl}" for i, lbl in enumerate(case["predictors"], start=1)
        )
        doc.add_paragraph()
        pp = doc.add_paragraph()
        pp.add_run(f"Beschikbare predictors:  {predictor_list}").font.size = Pt(10)

        doc.add_paragraph()
        instr = doc.add_paragraph()
        instr.add_run(
            "Opdracht: rangschik ALLE 5 predictors van hoogste naar laagste behandelprioriteit."
        ).bold = True
        instr.runs[0].font.size = Pt(10)
        for i in range(1, 6):
            _add_fillline(doc, f"Prioriteit {i}:", 55)

    _add_heading(doc, "4  Deel 3: Prioritering van behandeldoelen", level=1)
    _add_shaded_para(
        doc,
        "Rangschik de 5 standaardpredictoren van hoogste naar laagste klinische prioriteit als "
        "behandeldoel. Gebruik hiervoor de 21-daagse monitoring en het bipartiet netwerk "
        "(rood = sterke relatie, blauw = matige relatie). Zie de bijgevoegde PDF voor het "
        "kleurrijke netwerkdiagram.",
        shade_hex="FEF3C7",
    )
    _page_break(doc)
    word_part3_case(case_a)
    _page_break(doc)
    word_part3_case(case_b)
    _page_break(doc)

    # ── DEEL 4 ─────────────────────────────────────────────────────────────
    def word_part4_case(case: dict) -> None:
        _add_heading(doc, f"Casus {case['label']}: Deel 4", level=2)
        ctx = doc.add_paragraph()
        ctx.add_run("Gestandaardiseerde behandeldoelen uit Deel 3:").bold = True
        ctx.runs[0].font.size = Pt(10)
        for lbl, rationale in case["treatment_targets"]:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{lbl}  ").bold = True
            p.runs[0].font.size = Pt(9)
            p.add_run(f"({rationale})").font.size = Pt(9)
        doc.add_paragraph()
        instr = doc.add_paragraph()
        instr.add_run(
            "Opdracht: selecteer EXACT 5 EMA-items die het best passen als volgende subpredictoren. "
            "Vink de 5 geselecteerde items aan."
        ).bold = True
        instr.runs[0].font.size = Pt(10)
        doc.add_paragraph()
        for idx, item in enumerate(case["bfs_items"], start=1):
            _add_checkbox_item(doc, idx, item)
        doc.add_paragraph()
        total = doc.add_paragraph()
        total.add_run("Totaal geselecteerd: _____ / 5").bold = True
        total.runs[0].font.size = Pt(10)

    _add_heading(doc, "5  Deel 4: Verfijnd observatiemodel via breadth-first update-logica", level=1)
    _add_shaded_para(
        doc,
        "Selecteer per casus EXACT 5 EMA-items uit een lijst van 20. "
        "Start vanuit de gestandaardiseerde behandeldoelen en kies de 5 dagelijkse EMA-items "
        "die daar het best op aansluiten als directe subpredictoren. "
        "Verkies klinisch relevante breedte boven irrelevante of perifere items.",
        shade_hex="EDE9FE",
    )
    _page_break(doc)
    word_part4_case(case_a)
    _page_break(doc)
    word_part4_case(case_b)
    _page_break(doc)

    # ── DEEL 5 ─────────────────────────────────────────────────────────────
    def word_part5_case(case: dict) -> None:
        _add_heading(doc, f"Casus {case['label']}: Deel 5", level=2)
        ctx_data = [
            ("Primair probleem", case["p5_challenge"]),
            ("Behandeldoel", case["p5_target"]),
            ("Voornaamste barriere", case["p5_barrier"]),
            ("Copingstrategie", case["p5_coping"]),
        ]
        ctx_tbl = doc.add_table(rows=len(ctx_data), cols=2)
        ctx_tbl.style = "Table Grid"
        for row_i, (k, v) in enumerate(ctx_data):
            cells = ctx_tbl.rows[row_i].cells
            cells[0].width = Cm(4)
            cells[0].paragraphs[0].add_run(k).bold = True
            cells[0].paragraphs[0].runs[0].font.size = Pt(9)
            cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x04, 0x78, 0x57)
            cells[1].paragraphs[0].add_run(v).font.size = Pt(9)

        doc.add_paragraph()
        instr = doc.add_paragraph()
        instr.add_run(
            "Opdracht: schrijf hieronder de mobiele coachingsboodschap voor deze casus. "
            "Formuleer alsof de tekst morgen rechtstreeks op de smartphone van de persoon verschijnt."
        ).bold = True
        instr.runs[0].font.size = Pt(10)
        doc.add_paragraph()
        for _ in range(6):
            p = doc.add_paragraph("_" * 95)
            p.paragraph_format.space_after = Pt(4)
            p.runs[0].font.size = Pt(10)
            p.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    _add_heading(doc, "6  Deel 5: Mobiele coachingsboodschap", level=1)
    _add_shaded_para(
        doc,
        "Schrijf een korte, patientgerichte coachingsboodschap die rechtstreeks in de mobiele "
        "applicatie kan verschijnen. Gebruik het primaire probleem, het behandeldoel, de "
        "voornaamste barriere en de aangegeven copingstrategie. "
        "De boodschap moet compact, warm en professioneel zijn en een concrete eerste stap bevatten.\n\n"
        "Werkvoorbeeld (niet gerelateerd aan een studiecasus):\n"
        "\"Na een zware shift voelt rust nemen logisch, maar net dat eerste kleine beweegmoment "
        "kan je avond helpen ontladen. Trek vanavond meteen na thuiskomst je schoenen aan en "
        "wandel 10 minuten buiten -- dat ene blokje. Zo maak je de stap haalbaar.\"",
        shade_hex="FFF0F0",
    )
    _page_break(doc)
    word_part5_case(case_a)
    _page_break(doc)
    word_part5_case(case_b)
    _page_break(doc)

    # ── AFRONDING ──────────────────────────────────────────────────────────
    _add_heading(doc, "7  Afronding en terugbezorging", level=1)
    p = doc.add_paragraph(
        f"Dank u voor het invullen van alle vijf delen voor uw twee toegewezen casussen "
        f"({case_a['label']} en {case_b['label']})."
    )
    p.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    cl_head = doc.add_paragraph()
    cl_head.add_run("Checklist voor terugbezorging -- gelieve alle vakjes af te vinken:").bold = True
    cl_head.runs[0].font.size = Pt(10)

    checklist = [
        f"Ik heb in Deel 1 voor beide casussen ({case_a['label']} en {case_b['label']}) criteriumlabels ingevuld.",
        f"Ik heb in Deel 2 voor beide casussen predictorlabels ingevuld.",
        f"Ik heb in Deel 3 voor beide casussen alle 5 predictors gerangschikt.",
        f"Ik heb in Deel 4 voor beide casussen exact 5 EMA-items geselecteerd.",
        f"Ik heb in Deel 5 voor beide casussen een mobiele coachingsboodschap geschreven.",
        "Mijn antwoorden weerspiegelen mijn eigen klinische oordeel zonder gebruik van generatieve AI of andere externe hulp.",
        "Ik begrijp dat mijn antwoorden geanonimiseerd worden verwerkt en uitsluitend worden gebruikt binnen deze masterproefstudie.",
        "Ik begrijp dat ik mij op elk moment kan terugtrekken door contact op te nemen met de onderzoeker.",
    ]
    for item in checklist:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.3)
        p.paragraph_format.space_after = Pt(3)
        p.add_run(f"\u25a1   {item}").font.size = Pt(10)

    doc.add_paragraph()

    ret_head = doc.add_paragraph()
    ret_head.add_run("Terugbezorging:").bold = True
    ret_head.runs[0].font.size = Pt(10)

    ret = doc.add_paragraph(
        f"Bezorg dit ingevulde document als bijlage via e-mail:\n"
        f"  Aan:         stijn.vanseveren@ugent.be\n"
        f"  Onderwerp:   PHOENIX-PRE-{hcp_code}"
    )
    ret.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # Informed consent / anonymization final notice
    _add_shaded_para(
        doc,
        "VERTROUWELIJKHEID EN VERWERKING\n\n"
        "Na ontvangst worden uw antwoorden geanonimiseerd en opgenomen in het "
        "expertreferentiecorpus voor de latere blind evaluatie van het PHOENIX-systeem. "
        "Uw gegevens worden op geen enkel moment gekoppeld aan uw naam of identiteit in "
        "publieke of wetenschappelijke rapportage.\n\n"
        "Door dit document in te dienen bevestigt u vrijwillig in te stemmen met deelname "
        "aan deze studie. Uw inzending geldt als geïnformeerde toestemming (informed consent).\n\n"
        "Hartelijk dank voor uw bijdrage aan dit onderzoek.",
        shade_hex="DBEAFE",
    )

    return doc


def write_word_document(hcp_num: int) -> None:
    out_dir = BASE / f"HCP_{hcp_num}"
    out_dir.mkdir(parents=True, exist_ok=True)
    docx_path = out_dir / "main.docx"
    doc = build_word_document(hcp_num)
    doc.save(str(docx_path))
    print(f"[v] Word-document: HCP_{hcp_num}/main.docx")


if __name__ == "__main__":
    for idx in range(1, 6):
        write_and_compile(idx)
        write_word_document(idx)
    print("\nAlle 5 HCP-bundels zijn gegenereerd (PDF + Word).")
